"""Unit tests for spec gate orchestration and CLI command."""

from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document

from core.spec.spec_gate import evaluate_spec_gate, render_spec_gate_summary, run_spec_gate
from core.spec.spec_models import Finding
from etl.cli_spec import main


def _write_docx(path: Path, headings: list[str]) -> Path:
    """Create DOCX fixture with provided heading list."""
    doc = Document()
    for heading in headings:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(f"Texto de suporte: {heading}")
    doc.save(path)
    return path


def _write_mapping(path: Path, requirements: list[dict], requirement_ids: list[str]) -> Path:
    """Write mapping YAML fixture from requirement payload."""
    payload = {
        "schema_version": "1.0",
        "spec_source": "docs/analises/eventos/tamo_junto_2025/planning/00_docx_as_spec.md",
        "requirements": requirements,
        "marts": [
            {
                "mart_name": "mart_report_show_day_summary",
                "grain": "evento_dia_sessao",
                "description": "Resumo por sessao",
                "requirement_ids": requirement_ids,
                "output_fields": [
                    {
                        "field": "coverage_status",
                        "source": "public.mart_report_show_day_summary.coverage_status",
                        "rule": "Regra de cobertura",
                    }
                ],
            }
        ],
    }
    path.write_text(yaml.safe_dump(payload, sort_keys=False, allow_unicode=True), encoding="utf-8")
    return path


def _requirement(
    *,
    requirement_id: str,
    item_docx: str,
    target_field: str,
    calculation_rule: str,
) -> dict:
    """Build one requirement payload used by mapping fixtures."""
    return {
        "requirement_id": requirement_id,
        "item_docx": item_docx,
        "target_field": target_field,
        "calculation_rule": calculation_rule,
        "sources": [{"source_id": f"src_{requirement_id.lower()}", "location": "pagina 1 / tabela X"}],
        "validations": [
            {
                "rule_id": f"val_{requirement_id.lower()}",
                "rule_type": "reconciliation",
                "severity": "error",
                "expression": "1 = 1",
            }
        ],
    }


def test_run_spec_gate_returns_zero_when_no_error_findings(tmp_path: Path) -> None:
    """Gate returns zero when checks and mapping coverage have no errors."""
    docx_path = _write_docx(
        tmp_path / "ok.docx",
        headings=[
            "Publico do evento (controle de acesso - entradas validadas)",
            "Shows por dia (12/12, 13/12, 14/12)",
        ],
    )
    mapping_path = _write_mapping(
        tmp_path / "ok.yml",
        requirements=[
            _requirement(
                requirement_id="REQ-PUBLICO",
                item_docx="Publico do evento (controle de acesso - entradas validadas)",
                target_field="attendance_access_control.entries_validated",
                calculation_rule="SUM(entries_validated)",
            ),
            _requirement(
                requirement_id="REQ-SHOWS",
                item_docx="Shows por dia (12/12, 13/12, 14/12)",
                target_field="mart_report_show_day_summary.coverage_status",
                calculation_rule="CASE WHEN has_access_control THEN 'OK' ELSE 'GAP' END",
            ),
        ],
        requirement_ids=["REQ-PUBLICO", "REQ-SHOWS"],
    )

    rc = run_spec_gate(
        docx_path,
        mapping_path,
        required_sections=[
            "Publico do evento (controle de acesso - entradas validadas)",
            "Shows por dia (12/12, 13/12, 14/12)",
        ],
    )

    assert rc == 0


def test_evaluate_spec_gate_uses_policy_required_sections_by_default(tmp_path: Path) -> None:
    """Gate uses policy required sections when caller does not pass a custom list."""
    docx_path = _write_docx(
        tmp_path / "default_required.docx",
        headings=["Publico do evento (controle de acesso - entradas validadas)"],
    )
    mapping_path = _write_mapping(
        tmp_path / "default_required.yml",
        requirements=[
            _requirement(
                requirement_id="REQ-PUBLICO",
                item_docx="Publico do evento (controle de acesso - entradas validadas)",
                target_field="attendance_access_control.entries_validated",
                calculation_rule="SUM(entries_validated)",
            ),
            _requirement(
                requirement_id="REQ-SHOWS",
                item_docx="Shows por dia (12/12, 13/12, 14/12)",
                target_field="mart_report_show_day_summary.coverage_status",
                calculation_rule="CASE WHEN has_access_control THEN 'OK' ELSE 'GAP' END",
            ),
        ],
        requirement_ids=["REQ-PUBLICO", "REQ-SHOWS"],
    )

    rc, findings = evaluate_spec_gate(docx_path, mapping_path)

    assert rc == 1
    assert any(f.code == "SPEC_REQUIRED_SECTION_MISSING" for f in findings)


def test_evaluate_spec_gate_returns_error_findings_when_spec_is_incomplete(tmp_path: Path) -> None:
    """Gate evaluation reports actionable errors for missing section and mapping coverage."""
    docx_path = _write_docx(
        tmp_path / "incomplete.docx",
        headings=["Publico do evento (controle de acesso - entradas validadas)"],
    )
    mapping_path = _write_mapping(
        tmp_path / "incomplete.yml",
        requirements=[
            _requirement(
                requirement_id="REQ-PUBLICO",
                item_docx="Publico do evento (controle de acesso - entradas validadas)",
                target_field="attendance_access_control.entries_validated",
                calculation_rule="SUM(entries_validated)",
            )
        ],
        requirement_ids=["REQ-PUBLICO"],
    )

    rc, findings = evaluate_spec_gate(
        docx_path,
        mapping_path,
        required_sections=["Contexto do evento"],
    )
    codes = {finding.code for finding in findings}

    assert rc == 1
    assert "SPEC_REQUIRED_SECTION_MISSING" in codes
    assert "CHECKLIST_ITEM_WITHOUT_MAPPING" in codes


def test_render_spec_gate_summary_is_actionable() -> None:
    """Gate summary includes status, counts and concrete remediation instructions."""
    summary = render_spec_gate_summary(
        [
            Finding(
                severity="error",
                code="CHECKLIST_ITEM_WITHOUT_MAPPING",
                item_docx="Shows por dia (12/12, 13/12, 14/12)",
                message="Item do checklist nao possui mapping associado.",
                reference="Secao Shows",
            ),
            Finding(
                severity="warning",
                code="SPEC_SECTION_AMBIGUOUS_TITLE",
                item_docx="Publico do evento",
                message="Titulos ambiguos detectados para o mesmo item normalizado.",
                reference="Publico do evento; Publico do Evento",
            ),
        ]
    )

    assert "Status: FAIL" in summary
    assert "CHECKLIST_ITEM_WITHOUT_MAPPING" in summary
    assert "como corrigir" in summary


def test_cli_spec_gate_returns_non_zero_and_prints_summary(tmp_path: Path, capsys) -> None:
    """CLI command returns non-zero and prints actionable summary on gate failure."""
    docx_path = _write_docx(tmp_path / "cli.docx", headings=["Publico do evento"])
    mapping_path = _write_mapping(
        tmp_path / "cli.yml",
        requirements=[
            _requirement(
                requirement_id="REQ-CLI",
                item_docx="Publico do evento",
                target_field="attendance_access_control.entries_validated",
                calculation_rule="SUM(entries_validated)",
            )
        ],
        requirement_ids=["REQ-CLI"],
    )

    rc = main(
        [
            "spec:gate",
            "--docx",
            str(docx_path),
            "--mapping",
            str(mapping_path),
            "--required-section",
            "Contexto do evento",
        ]
    )
    captured = capsys.readouterr()

    assert rc == 1
    assert "Spec Gate Summary" in captured.out
    assert "Status: FAIL" in captured.out
