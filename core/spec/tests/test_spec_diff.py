"""Unit tests for DOCX spec diff utilities and CLI integration."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from core.spec.spec_diff import diff_specs, render_diff_json, render_diff_markdown
from etl.cli_spec import _build_spec, main


def _create_docx_old(path: Path) -> Path:
    """Create baseline DOCX fixture for diff tests."""
    doc = Document()
    doc.add_heading("Contexto do evento", level=1)
    doc.add_paragraph("Figura 1 - Entradas validadas por dia")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Metrica"
    tbl.rows[0].cells[1].text = "Valor"
    tbl.rows[1].cells[0].text = "Entradas validadas"
    tbl.rows[1].cells[1].text = "100"
    doc.save(path)
    return path


def _create_docx_new(path: Path) -> Path:
    """Create candidate DOCX fixture with intentional structural changes."""
    doc = Document()
    doc.add_heading("Contexto do evento", level=1)
    doc.add_heading("Shows por dia", level=1)
    doc.add_paragraph("Figura 2 - Opt-in aceitos por dia")
    tbl = doc.add_table(rows=2, cols=3)
    tbl.rows[0].cells[0].text = "Metrica"
    tbl.rows[0].cells[1].text = "Dia"
    tbl.rows[0].cells[2].text = "Valor"
    tbl.rows[1].cells[0].text = "Opt-in aceitos"
    tbl.rows[1].cells[1].text = "12/12"
    tbl.rows[1].cells[2].text = "120"
    doc.save(path)
    return path


def test_diff_specs_detects_added_and_removed_sections_figures_tables(tmp_path: Path) -> None:
    """Diff report identifies structural additions and removals."""
    old_docx = _create_docx_old(tmp_path / "old.docx")
    new_docx = _create_docx_new(tmp_path / "new.docx")

    report = diff_specs(_build_spec(old_docx), _build_spec(new_docx))

    assert "Shows por dia" in report.sections_added
    assert "Figura 2 - Opt-in aceitos por dia" in report.figures_added
    assert "Figura 1 - Entradas validadas por dia" in report.figures_removed
    assert any("cols=3" in table for table in report.tables_added)
    assert any("cols=2" in table for table in report.tables_removed)


def test_spec_diff_renderers_generate_markdown_and_json(tmp_path: Path) -> None:
    """Markdown and JSON renders expose diff report data consistently."""
    old_docx = _create_docx_old(tmp_path / "old.docx")
    new_docx = _create_docx_new(tmp_path / "new.docx")
    report = diff_specs(_build_spec(old_docx), _build_spec(new_docx))

    markdown = render_diff_markdown(report)
    as_json = render_diff_json(report)
    payload = json.loads(as_json)

    assert "## Secoes" in markdown
    assert "| adicionado | Shows por dia |" in markdown
    assert "old_source" in payload
    assert "figures" in payload
    assert payload["sections"]["added"] == report.sections_added


def test_cli_spec_diff_writes_markdown_and_json_files(tmp_path: Path) -> None:
    """CLI `spec:diff` writes both markdown and json outputs."""
    old_docx = _create_docx_old(tmp_path / "old.docx")
    new_docx = _create_docx_new(tmp_path / "new.docx")
    out_md = tmp_path / "spec_diff.md"
    out_json = tmp_path / "spec_diff.json"

    rc = main(
        [
            "spec:diff",
            "--old-docx",
            str(old_docx),
            "--new-docx",
            str(new_docx),
            "--out-md",
            str(out_md),
            "--out-json",
            str(out_json),
        ]
    )

    assert rc == 0
    assert out_md.exists()
    assert out_json.exists()
    assert "Spec Diff Summary" in out_md.read_text(encoding="utf-8")
    payload = json.loads(out_json.read_text(encoding="utf-8"))
    assert "sections" in payload
    assert "tables" in payload
