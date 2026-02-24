"""DOCX inventory extractor for figures and tables.

This focuses on "structure" (captions, headers), not on extracting the raw
numbers from figures (which would require additional parsing/OCR and is out of
scope for the spec extractor).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, Optional, Sequence

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .docx_sections import _heading_level, find_nearest_section_title
from .spec_models import DocxFigure, DocxSection, DocxTable


_FIGURE_RE = re.compile(r"^\s*Figura\b", flags=re.IGNORECASE)


def _iter_block_items(doc: DocxDocument) -> Iterable[Paragraph | Table]:
    """Yield paragraphs and tables in document order."""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _is_figure_caption(paragraph: Paragraph) -> bool:
    """Return True when a paragraph represents a figure caption.

    Args:
        paragraph: A DOCX paragraph.

    Returns:
        True for paragraphs styled as caption or starting with "Figura".
    """
    text = (paragraph.text or "").strip()
    if not text:
        return False
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name and style_name.strip().lower() == "caption" and _FIGURE_RE.search(text):
        return True
    return bool(_FIGURE_RE.search(text))


def extract_figures(
    paragraphs: Sequence[Paragraph],
    sections: Sequence[DocxSection],
) -> List[DocxFigure]:
    """Extract figure captions and associate each item to the nearest section.

    Args:
        paragraphs: Paragraph sequence in document order.
        sections: Extracted sections in document order.

    Returns:
        Extracted figure inventory with caption text and paragraph position.

    Raises:
        ValueError: If sections are not in ascending paragraph order.
    """
    if any(
        sections[i].start_paragraph_index > sections[i + 1].start_paragraph_index
        for i in range(len(sections) - 1)
    ):
        raise ValueError("Sections must be ordered by start_paragraph_index.")

    figures: List[DocxFigure] = []
    for paragraph_index, paragraph in enumerate(paragraphs):
        if not _is_figure_caption(paragraph):
            continue
        figures.append(
            DocxFigure(
                caption=(paragraph.text or "").strip(),
                paragraph_index=paragraph_index,
                section_title=find_nearest_section_title(paragraph_index, sections),
            )
        )
    return figures


def extract_tables(doc: DocxDocument) -> List[DocxTable]:
    """Extract table headers and associate each table to the closest heading above.

    Args:
        doc: Loaded DOCX document.

    Returns:
        Extracted table inventory with headers and shape.
    """
    tables: List[DocxTable] = []
    current_section_title: Optional[str] = None

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = (block.text or "").strip()
            if not text:
                continue
            style_name = block.style.name if block.style else ""
            if _heading_level(style_name) is not None:
                current_section_title = text
            continue

        # Table
        t: Table = block
        if not t.rows:
            continue
        header_cells = [c.text.strip() for c in t.rows[0].cells]
        n_cols = len(t.rows[0].cells)
        n_rows = len(t.rows)
        tables.append(
            DocxTable(
                table_index=len(tables) + 1,
                n_rows=n_rows,
                n_cols=n_cols,
                header_cells=header_cells,
                section_title=current_section_title,
            )
        )

    return tables


def extract_figures_from_docx(docx_path: str | Path, sections: Sequence[DocxSection]) -> List[DocxFigure]:
    """Load a DOCX and delegate figure extraction.

    Args:
        docx_path: Path to a `.docx` file.
        sections: Extracted sections in document order.

    Returns:
        Figure inventory extracted from the DOCX.

    Raises:
        FileNotFoundError: If the path does not exist.
        ValueError: If sections are not ordered by paragraph index.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    doc = Document(path)
    return extract_figures(doc.paragraphs, sections)


def extract_tables_from_docx(docx_path: str | Path) -> List[DocxTable]:
    """Load a DOCX and delegate table extraction.

    Args:
        docx_path: Path to a `.docx` file.

    Returns:
        Table inventory extracted from the DOCX.

    Raises:
        FileNotFoundError: If the path does not exist.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    doc = Document(path)
    return extract_tables(doc)
