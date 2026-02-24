"""Unit tests for Word report scaffold renderer and CLI."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from reports.word import WordReportRenderer
from reports.word import cli_report


def _build_min_template(path: Path) -> Path:
    """Create a minimal DOCX template used by scaffold tests."""

    document = Document()
    document.add_heading("Template TMJ 2025", level=1)
    document.add_paragraph("Conteudo base do template.")
    document.save(path)
    return path


def test_word_report_renderer_generates_output_docx_from_template(tmp_path: Path) -> None:
    """Renderer should create output DOCX and append scaffold execution marker."""

    template_path = _build_min_template(tmp_path / "template_min.docx")
    output_path = tmp_path / "out" / "report_out.docx"

    renderer = WordReportRenderer(
        event_id=2025,
        template_path=template_path,
        output_path=output_path,
    )
    result = renderer.render()

    assert result.event_id == 2025
    assert result.template_path == template_path
    assert result.output_path == output_path
    assert output_path.exists()

    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Template TMJ 2025" in all_text
    assert "[NPBB scaffold] Report generated automatically." in all_text
    assert "event_id=2025" in all_text


def test_report_render_cli_generates_docx_and_returns_zero(tmp_path: Path) -> None:
    """CLI `report:render` should produce DOCX output and return exit code zero."""

    template_path = _build_min_template(tmp_path / "template_cli.docx")
    output_path = tmp_path / "generated" / "report_cli.docx"

    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "77",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
        ]
    )

    assert rc == 0
    assert output_path.exists()


def test_report_render_cli_returns_one_for_missing_template(tmp_path: Path) -> None:
    """CLI should return non-zero when template path is missing."""

    output_path = tmp_path / "generated" / "report_missing.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "10",
            "--template-path",
            str(tmp_path / "nao_existe.docx"),
            "--output-path",
            str(output_path),
        ]
    )

    assert rc == 1
    assert not output_path.exists()


def test_word_report_renderer_processes_text_placeholder_payload(tmp_path: Path) -> None:
    """Renderer should resolve mapped text placeholder with readable bullet lines."""

    template_path = tmp_path / "template_text.docx"
    document = Document()
    document.add_heading("Resumo", level=1)
    document.add_paragraph("{{FONTES__SUMMARY__TEXT}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_text_out.docx"
    renderer = WordReportRenderer(
        event_id=2026,
        template_path=template_path,
        output_path=output_path,
        text_payload_by_placeholder={
            "FONTES__SUMMARY__TEXT": [
                "Fonte de acesso validada para show 12/12",
                "Fonte de vendas ausente para show 14/12 (GAP formal)",
            ]
        },
    )
    renderer.render()

    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{FONTES__SUMMARY__TEXT}}" not in all_text
    assert "Fonte de acesso validada para show 12/12" in all_text
    assert "Fonte de vendas ausente para show 14/12 (GAP formal)" in all_text


def test_word_report_renderer_rejects_payload_for_non_text_placeholder(tmp_path: Path) -> None:
    """Renderer should fail when payload targets placeholder not mapped as text."""

    template_path = _build_min_template(tmp_path / "template_invalid_payload.docx")
    output_path = tmp_path / "out" / "report_invalid_payload.docx"
    renderer = WordReportRenderer(
        event_id=2026,
        template_path=template_path,
        output_path=output_path,
        text_payload_by_placeholder={"PUBLICO__ATTENDANCE__TABLE": "payload invalido"},
    )

    with pytest.raises(ValueError, match="nao mapeado\\(s\\) como text"):
        renderer.render()


def test_report_render_cli_fails_for_unknown_placeholder_in_template(tmp_path: Path) -> None:
    """CLI should fail when template has placeholder without mapping contract."""

    template_path = tmp_path / "template_unknown.docx"
    document = Document()
    document.add_paragraph("{{PLACEHOLDER__DESCONHECIDO__TEXT}}")
    document.save(template_path)

    output_path = tmp_path / "generated" / "report_unknown.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "88",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
        ]
    )

    assert rc == 1
    assert not output_path.exists()


def test_report_render_cli_fails_for_unresolved_placeholder_after_render(tmp_path: Path) -> None:
    """CLI should fail when mapped placeholder remains unresolved in output DOCX."""

    template_path = tmp_path / "template_unresolved.docx"
    document = Document()
    document.add_paragraph("{{FONTES__SUMMARY__TEXT}}")
    document.save(template_path)

    output_path = tmp_path / "generated" / "report_unresolved.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "89",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
        ]
    )

    assert rc == 1
    assert output_path.exists()


def test_report_render_cli_passes_when_text_placeholders_are_resolved(tmp_path: Path) -> None:
    """CLI should pass and generate final DOCX when text placeholders are resolved."""

    template_path = tmp_path / "template_resolved.docx"
    document = Document()
    document.add_paragraph("{{FONTES__SUMMARY__TEXT}}")
    document.save(template_path)

    payload_path = tmp_path / "payload.json"
    payload_path.write_text(
        json.dumps(
            {
                "FONTES__SUMMARY__TEXT": [
                    "Fonte A: controle de acesso ok",
                    "Fonte B: vendas com GAP formal",
                ]
            }
        ),
        encoding="utf-8",
    )

    output_path = tmp_path / "generated" / "report_resolved.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "90",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--text-payload-json",
            str(payload_path),
        ]
    )

    assert rc == 0
    assert output_path.exists()
    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{FONTES__SUMMARY__TEXT}}" not in all_text


def test_report_render_cli_runs_preflight_and_renders_show_coverage(tmp_path: Path, monkeypatch) -> None:
    """CLI should run DQ+coverage preflight and pass artifacts to report renderer."""

    template_path = tmp_path / "template_preflight.docx"
    document = Document()
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.add_paragraph("{{SHOW__COVERAGE__TABLE}}")
    document.save(template_path)

    dq_out = tmp_path / "artifacts" / "dq_report.json"
    coverage_out = tmp_path / "artifacts" / "coverage_report.json"

    def _fake_run_dq_checks(**kwargs):  # noqa: ANN001
        payload = {
            "summary": {"error": 0, "warning": 0, "info": 1},
            "sections": {
                "errors": [],
                "warnings": [],
                "inconsistencies": [],
                "gaps_by_dataset": {},
            },
        }
        kwargs["out_json"].parent.mkdir(parents=True, exist_ok=True)
        kwargs["out_json"].write_text(json.dumps(payload), encoding="utf-8")
        return 0

    def _fake_run_show_coverage_report(*, event_id, out_json, contract_path):  # noqa: ANN001
        assert event_id == 2025
        _ = contract_path
        payload = {
            "generated_at": "2026-02-11T18:30:00+00:00",
            "event_id": 2025,
            "contract_version": 5,
            "status": "partial",
            "summary": {
                "total_sessions": 1,
                "ok_sessions": 0,
                "partial_sessions": 1,
                "gap_sessions": 0,
                "total_missing_inputs": 1,
                "gap_missing_inputs": 0,
                "partial_missing_inputs": 1,
            },
            "sessions": [
                {
                    "session_id": 11,
                    "session_key": "TMJ2025_20251212_SHOW",
                    "session_date": "2025-12-12",
                    "session_type": "noturno_show",
                    "status": "partial",
                    "observed_datasets": ["access_control"],
                    "partial_datasets": ["optin"],
                    "missing_datasets": ["ticket_sales"],
                    "missing_inputs": [
                        {
                            "dataset": "ticket_sales",
                            "reason_code": "missing_in_catalog",
                            "reason": "Nenhuma fonte encontrada.",
                        }
                    ],
                }
            ],
            "missing_inputs": [],
        }
        out_json.parent.mkdir(parents=True, exist_ok=True)
        out_json.write_text(json.dumps(payload), encoding="utf-8")
        return payload

    class _Agenda:
        version = 7
        source_path = Path("fake_agenda_master.yml")

    monkeypatch.setattr(cli_report, "run_dq_checks", _fake_run_dq_checks)
    monkeypatch.setattr(cli_report, "run_show_coverage_report", _fake_run_show_coverage_report)
    monkeypatch.setattr(cli_report, "load_agenda_master", lambda *_args, **_kwargs: _Agenda())

    output_path = tmp_path / "generated" / "report_preflight.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "2025",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--run-preflight",
            "--ingestion-id",
            "42",
            "--dq-out-json",
            str(dq_out),
            "--coverage-out-json",
            str(coverage_out),
        ]
    )

    assert rc == 0
    assert output_path.exists()
    assert dq_out.exists()
    assert coverage_out.exists()

    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{GAPS__SUMMARY__TEXT}}" not in all_text
    assert "{{SHOW__COVERAGE__TABLE}}" not in all_text
    assert "Legenda de status da cobertura de shows por dia:" in all_text
    assert len(rendered.tables) == 1

    manifest_path = output_path.parent / "report_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    governance = manifest["coverage_governance"]
    assert governance["agenda_version"] == 7
    assert governance["coverage_contract_version"] == 5
    assert governance["dq_report_path"] == str(dq_out)
    assert governance["coverage_report_path"] == str(coverage_out)


def test_report_render_cli_partial_on_critical_show_gap_includes_coverage_and_gap_sections(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Critical show gap should force partial and render both coverage and GAP sections."""

    template_path = tmp_path / "template_show_gap_partial.docx"
    document = Document()
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.add_paragraph("{{SHOW__COVERAGE__TABLE}}")
    document.save(template_path)

    dq_out = tmp_path / "artifacts" / "dq_report_show_gap.json"
    coverage_out = tmp_path / "artifacts" / "coverage_report_show_gap.json"

    def _fake_run_dq_checks(**kwargs):  # noqa: ANN001
        payload = {
            "summary": {"error": 0, "warning": 0, "info": 0},
            "sections": {
                "errors": [],
                "warnings": [],
                "inconsistencies": [],
                "gaps_by_dataset": {},
            },
        }
        kwargs["out_json"].parent.mkdir(parents=True, exist_ok=True)
        kwargs["out_json"].write_text(json.dumps(payload), encoding="utf-8")
        return 0

    def _fake_run_show_coverage_report(*, event_id, out_json, contract_path):  # noqa: ANN001
        assert event_id == 2025
        _ = contract_path
        payload = {
            "generated_at": "2026-02-11T19:00:00+00:00",
            "event_id": 2025,
            "contract_version": 6,
            "status": "gap",
            "summary": {
                "total_sessions": 1,
                "ok_sessions": 0,
                "partial_sessions": 0,
                "gap_sessions": 1,
                "total_missing_inputs": 1,
                "gap_missing_inputs": 1,
                "partial_missing_inputs": 0,
            },
            "sessions": [
                {
                    "session_id": 12,
                    "session_key": "TMJ2025_20251214_SHOW",
                    "session_date": "2025-12-14",
                    "session_type": "NOTURNO_SHOW",
                    "status": "gap",
                    "observed_datasets": ["optin"],
                    "partial_datasets": [],
                    "missing_datasets": ["access_control", "ticket_sales"],
                    "missing_inputs": [
                        {
                            "dataset": "access_control",
                            "reason_code": "missing_in_catalog",
                            "reason": "Nenhuma fonte correspondente encontrada.",
                        }
                    ],
                }
            ],
            "missing_inputs": [],
        }
        out_json.parent.mkdir(parents=True, exist_ok=True)
        out_json.write_text(json.dumps(payload), encoding="utf-8")
        return payload

    class _Agenda:
        version = 8
        source_path = Path("fake_agenda_master_show_gap.yml")

    gate_policy_path = tmp_path / "gate_show_gap_partial.yml"
    gate_policy_path.write_text(
        "\n".join(
            [
                "version: 1",
                "block:",
                "  enabled: true",
                "  min_error_findings: 999",
                "  critical_statuses: [GAP]",
                "  critical_tokens: [unused]",
                "partial:",
                "  enabled: true",
                "  statuses: [GAP, INCONSISTENTE]",
                "  min_findings: 1",
                "  banner_text: \"PARCIAL: show coverage critico\"",
                "output:",
                "  allow_partial_on_block: true",
                "show_coverage:",
                "  enabled: true",
                "  session_type: NOTURNO_SHOW",
                "  critical_dataset: access_control",
                "  critical_statuses: [GAP]",
                "  required_placeholders: [SHOW__COVERAGE__TABLE, GAPS__SUMMARY__TEXT]",
            ]
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(cli_report, "run_dq_checks", _fake_run_dq_checks)
    monkeypatch.setattr(cli_report, "run_show_coverage_report", _fake_run_show_coverage_report)
    monkeypatch.setattr(cli_report, "load_agenda_master", lambda *_args, **_kwargs: _Agenda())

    output_path = tmp_path / "generated" / "report_show_gap_partial.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "2025",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--run-preflight",
            "--ingestion-id",
            "99",
            "--dq-out-json",
            str(dq_out),
            "--coverage-out-json",
            str(coverage_out),
            "--gate-policy-path",
            str(gate_policy_path),
        ]
    )

    assert rc == 0
    assert output_path.exists()
    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[STATUS] PARCIAL: show coverage critico" in all_text
    assert "Legenda de status da cobertura de shows por dia:" in all_text
    assert "TMJ2025_20251214_SHOW_access_control" in all_text
    assert len(rendered.tables) == 1


def test_report_render_cli_preflight_requires_ingestion_id(tmp_path: Path) -> None:
    """CLI should fail when preflight is requested without ingestion-id."""

    template_path = _build_min_template(tmp_path / "template_preflight_missing_ingestion.docx")
    output_path = tmp_path / "generated" / "report_preflight_missing_ingestion.docx"

    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "2025",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--run-preflight",
        ]
    )

    assert rc == 1
    assert not output_path.exists()


def test_word_report_renderer_processes_table_placeholder_payload(tmp_path: Path) -> None:
    """Renderer should resolve mapped table placeholder with contract columns."""

    template_path = tmp_path / "template_table.docx"
    document = Document()
    document.add_heading("Publico", level=1)
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_table_out.docx"
    renderer = WordReportRenderer(
        event_id=2030,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": {
                "columns": ["sessao", "presentes", "validos"],
                "rows": [
                    {"sessao": "show_12", "presentes": 100, "validos": 120},
                    {"sessao": "show_13", "presentes": 150, "validos": 180},
                ],
            }
        },
    )
    renderer.render()

    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{PUBLICO__ATTENDANCE__TABLE}}" not in all_text
    assert len(rendered.tables) == 1
    table = rendered.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == ["sessao", "presentes", "validos"]
    assert [cell.text for cell in table.rows[1].cells] == ["show_12", "100", "120"]
    assert [cell.text for cell in table.rows[2].cells] == ["show_13", "150", "180"]


def test_word_report_renderer_blocks_table_payload_contract_failure(tmp_path: Path) -> None:
    """Renderer should fail when table payload misses required contract columns."""

    template_path = tmp_path / "template_table_invalid.docx"
    document = Document()
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.save(template_path)

    renderer = WordReportRenderer(
        event_id=2030,
        template_path=template_path,
        output_path=tmp_path / "out" / "report_table_invalid.docx",
        table_payload_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": {
                "columns": ["sessao", "presentes", "validos"],
                "rows": [{"sessao": "show_12", "presentes": 100}],
            }
        },
    )

    with pytest.raises(ValueError, match="Contrato de colunas violado"):
        renderer.render()


def test_word_report_renderer_rejects_payload_for_non_table_placeholder(tmp_path: Path) -> None:
    """Renderer should fail when table payload targets placeholder not mapped as table."""

    template_path = _build_min_template(tmp_path / "template_invalid_table_payload.docx")
    output_path = tmp_path / "out" / "report_invalid_table_payload.docx"
    renderer = WordReportRenderer(
        event_id=2031,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "FONTES__SUMMARY__TEXT": {
                "columns": ["coluna_a"],
                "rows": [["valor"]],
            }
        },
    )

    with pytest.raises(ValueError, match="nao mapeado\\(s\\) como table"):
        renderer.render()


def test_word_report_renderer_renders_lineage_block_for_table(tmp_path: Path) -> None:
    """Renderer should append lineage block for each rendered table placeholder."""

    template_path = tmp_path / "template_lineage_table.docx"
    document = Document()
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_lineage_table.docx"
    renderer = WordReportRenderer(
        event_id=2032,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": {
                "columns": ["sessao", "presentes", "validos"],
                "rows": [{"sessao": "show_12", "presentes": 100, "validos": 120}],
            }
        },
        table_lineage_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": [
                {
                    "source_id": "pdf_access_12",
                    "location": "page=3",
                    "evidence_text": "Tabela Controle de Acesso",
                }
            ]
        },
    )
    renderer.render()

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[Linhagem] table | placeholder=PUBLICO__ATTENDANCE__TABLE" in text
    assert "source_id=pdf_access_12" in text
    assert "location=page=3" in text
    assert "evidence_text=Tabela Controle de Acesso" in text


def test_word_report_renderer_renders_show_coverage_section_with_legend(tmp_path: Path) -> None:
    """Renderer should include show-coverage table and legend for status semantics."""

    template_path = tmp_path / "template_show_coverage.docx"
    document = Document()
    document.add_paragraph("{{SHOW__COVERAGE__TABLE}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_show_coverage.docx"
    renderer = WordReportRenderer(
        event_id=2033,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "SHOW__COVERAGE__TABLE": {
                "columns": [
                    "dia",
                    "sessao",
                    "status_access_control",
                    "status_optin",
                    "status_ticket_sales",
                    "observacoes",
                ],
                "rows": [
                    {
                        "dia": "2025-12-12",
                        "sessao": "Show 12/12",
                        "status_access_control": "ok",
                        "status_optin": "partial",
                        "status_ticket_sales": "gap",
                        "observacoes": "falta_vendidos_total",
                    }
                ],
            }
        },
    )
    renderer.render()

    rendered = Document(output_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{SHOW__COVERAGE__TABLE}}" not in all_text
    assert "Legenda de status da cobertura de shows por dia:" in all_text
    assert "OK: dataset presente e carregado para a sessao." in all_text
    assert "GAP: dataset ausente para a sessao esperada." in all_text
    assert "INCONSISTENTE: dataset presente com conflito entre fontes ou regras." in all_text
    assert len(rendered.tables) == 1
    table_header = [cell.text for cell in rendered.tables[0].rows[0].cells]
    assert table_header == [
        "dia",
        "sessao",
        "status_access_control",
        "status_optin",
        "status_ticket_sales",
        "observacoes",
    ]


def test_word_report_renderer_renders_gap_when_figure_lineage_is_missing(tmp_path: Path) -> None:
    """Renderer should write GAP lineage block when figure lineage is not provided."""

    from reports.charts import ChartSeries, ChartSpec, ChartType, render_chart_png

    chart_spec = ChartSpec(
        chart_type=ChartType.BAR,
        title="Presencas por dia",
        x_label="Dia",
        y_label="Presentes",
        series=(ChartSeries(name="Presencas", x_field="dia", y_field="presentes"),),
    )
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(
        render_chart_png(
            chart_spec,
            [
                {"dia": "2025-12-12", "presentes": 120},
                {"dia": "2025-12-13", "presentes": 180},
            ],
        )
    )

    template_path = tmp_path / "template_lineage_figure.docx"
    document = Document()
    document.add_paragraph("{{PRESALE__CURVES__FIGURE}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_lineage_figure.docx"
    renderer = WordReportRenderer(
        event_id=2033,
        template_path=template_path,
        output_path=output_path,
        figure_payload_by_placeholder={
            "PRESALE__CURVES__FIGURE": {
                "image_path": str(image_path),
                "caption": "Figura 2 - Curva de presencas",
            }
        },
    )
    renderer.render()

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[Linhagem] figure | placeholder=PRESALE__CURVES__FIGURE" in text
    assert "GAP: linhagem ausente" in text


def test_word_report_renderer_renders_gap_section_from_dq_report(tmp_path: Path) -> None:
    """Renderer should fill dedicated GAP placeholder from DQ report payload."""

    template_path = tmp_path / "template_gaps_section.docx"
    document = Document()
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.save(template_path)

    dq_report = {
        "sections": {
            "inconsistencies": [
                {
                    "dataset_id": "show_12_dimac",
                    "message": "Conflito em show_12",
                    "source_id": "SRC_DIMAC",
                    "location": "page=8",
                    "evidence_text": "Tabela de share",
                }
            ],
            "gaps_by_dataset": {
                "show_14_access_control": [
                    {
                        "message": "Falta controle de acesso em 2025-12-14",
                        "source_id": "SRC_ACCESS",
                        "location": "page=0",
                        "evidence_text": "Arquivo nao entregue",
                    }
                ]
            },
        }
    }

    output_path = tmp_path / "out" / "report_gaps_section.docx"
    renderer = WordReportRenderer(
        event_id=2034,
        template_path=template_path,
        output_path=output_path,
        dq_report=dq_report,
        gaps_placeholder_id="GAPS__SUMMARY__TEXT",
    )
    renderer.render()

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{GAPS__SUMMARY__TEXT}}" not in text
    assert "Resumo GAP/INCONSISTENTE" in text
    assert "show_12_dimac" in text
    assert "show_14_access_control" in text
    assert "12/12: INCONSISTENTE" in text
    assert "14/12: GAP" in text


def test_report_render_cli_blocks_generation_when_gate_policy_blocks(tmp_path: Path) -> None:
    """CLI should stop report generation when publication gate returns block."""

    template_path = tmp_path / "template_gate_block.docx"
    document = Document()
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.save(template_path)

    dq_report_path = tmp_path / "dq_block.json"
    dq_report_path.write_text(
        json.dumps(
            {
                "summary": {"error": 0},
                "sections": {
                    "gaps_by_dataset": {
                        "show_14_access_control": [
                            {"message": "Falta fonte oficial"}
                        ]
                    }
                },
            }
        ),
        encoding="utf-8",
    )
    gate_policy_path = tmp_path / "gate_block.yml"
    gate_policy_path.write_text(
        "\n".join(
            [
                "version: 1",
                "block:",
                "  enabled: true",
                "  min_error_findings: 0",
                "  critical_statuses: [GAP]",
                "  critical_tokens: [show_14]",
                "partial:",
                "  enabled: true",
                "  statuses: [GAP]",
                "  min_findings: 1",
                "  banner_text: PARCIAL",
                "output:",
                "  allow_partial_on_block: false",
            ]
        ),
        encoding="utf-8",
    )

    output_path = tmp_path / "out" / "report_gate_block.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "2035",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--dq-report-json",
            str(dq_report_path),
            "--gate-policy-path",
            str(gate_policy_path),
        ]
    )

    assert rc == 1
    assert not output_path.exists()


def test_report_render_cli_generates_partial_when_gate_allows_partial(tmp_path: Path) -> None:
    """CLI should generate DOCX with explicit partial status when policy allows it."""

    template_path = tmp_path / "template_gate_partial.docx"
    document = Document()
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.save(template_path)

    dq_report_path = tmp_path / "dq_partial.json"
    dq_report_path.write_text(
        json.dumps(
            {
                "summary": {"error": 0},
                "sections": {
                    "gaps_by_dataset": {
                        "show_12_access_control": [
                            {"message": "Sem fonte oficial"}
                        ]
                    }
                },
            }
        ),
        encoding="utf-8",
    )
    gate_policy_path = tmp_path / "gate_partial.yml"
    gate_policy_path.write_text(
        "\n".join(
            [
                "version: 1",
                "block:",
                "  enabled: true",
                "  min_error_findings: 0",
                "  critical_statuses: [GAP]",
                "  critical_tokens: [show_12]",
                "partial:",
                "  enabled: true",
                "  statuses: [GAP]",
                "  min_findings: 1",
                "  banner_text: \"PARCIAL: revisar gaps\"",
                "output:",
                "  allow_partial_on_block: true",
            ]
        ),
        encoding="utf-8",
    )

    output_path = tmp_path / "out" / "report_gate_partial.docx"
    rc = cli_report.main(
        [
            "report:render",
            "--event-id",
            "2036",
            "--template-path",
            str(template_path),
            "--output-path",
            str(output_path),
            "--dq-report-json",
            str(dq_report_path),
            "--gate-policy-path",
            str(gate_policy_path),
        ]
    )

    assert rc == 0
    assert output_path.exists()
    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[STATUS] PARCIAL: revisar gaps" in text
    assert "show_12_access_control" in text
