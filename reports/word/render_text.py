"""Text placeholder rendering helpers for Word report sections.

This module provides minimal text rendering utilities to replace placeholders
inside a DOCX document with readable text paragraphs or bullet lists.
"""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from typing import Any

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


class PlaceholderTextRenderError(ValueError):
    """Raised when text placeholder rendering fails."""


def placeholder_token(placeholder_id: str) -> str:
    """Build DOCX token string from placeholder id.

    Args:
        placeholder_id: Canonical placeholder id from mapping config.

    Returns:
        Token formatted as `{{PLACEHOLDER_ID}}`.
    """

    return f"{{{{{placeholder_id}}}}}"


def _iter_document_paragraphs(doc: DocxDocument) -> list[Paragraph]:
    """Collect all paragraphs from document body and table cells.

    Args:
        doc: Open DOCX document.

    Returns:
        Flat list with body and table-cell paragraphs.
    """

    paragraphs: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _normalize_text_payload(value: Any, *, placeholder_id: str) -> list[str]:
    """Normalize text payload into a non-empty list of strings.

    Args:
        value: Raw payload (`str` or `list[str]`).
        placeholder_id: Placeholder identifier for error context.

    Returns:
        Normalized list of text lines.

    Raises:
        PlaceholderTextRenderError: If payload is unsupported or empty.
    """

    if isinstance(value, str):
        clean = value.strip()
        if not clean:
            raise PlaceholderTextRenderError(
                f"Payload invalido para {placeholder_id}: texto vazio."
            )
        return [clean]

    if isinstance(value, Sequence) and not isinstance(value, (bytes, bytearray, str)):
        lines: list[str] = []
        for idx, item in enumerate(value):
            if not isinstance(item, str) or not item.strip():
                raise PlaceholderTextRenderError(
                    f"Payload invalido para {placeholder_id}: item {idx} deve ser texto nao vazio."
                )
            lines.append(item.strip())
        if not lines:
            raise PlaceholderTextRenderError(
                f"Payload invalido para {placeholder_id}: lista vazia."
            )
        return lines

    raise PlaceholderTextRenderError(
        f"Payload invalido para {placeholder_id}: esperado str ou list[str]."
    )


def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    """Insert a paragraph right after an existing paragraph.

    Args:
        anchor: Paragraph that will receive a new sibling after it.

    Returns:
        Newly created paragraph object.
    """

    new_node = OxmlElement("w:p")
    anchor._p.addnext(new_node)
    return Paragraph(new_node, anchor._parent)


def render_text_placeholder(
    doc: DocxDocument,
    placeholder_id: str,
    value: str | Sequence[str],
) -> int:
    """Render one text placeholder in a DOCX document.

    Rules:
    - placeholder token format is `{{PLACEHOLDER_ID}}`;
    - scalar text replaces token directly;
    - list payload renders first line in place and remaining lines as bullets.

    Args:
        doc: Open DOCX document.
        placeholder_id: Placeholder identifier to replace.
        value: Replacement payload (`str` or `list[str]`).

    Returns:
        Number of placeholder occurrences replaced.

    Raises:
        PlaceholderTextRenderError: If placeholder is missing or payload invalid.
    """

    token = placeholder_token(placeholder_id)
    lines = _normalize_text_payload(value, placeholder_id=placeholder_id)
    paragraphs = _iter_document_paragraphs(doc)

    matched = [paragraph for paragraph in paragraphs if token in paragraph.text]
    if not matched:
        raise PlaceholderTextRenderError(
            f"Placeholder de texto nao encontrado no DOCX: {placeholder_id}."
        )

    for paragraph in matched:
        original_text = paragraph.text
        if len(lines) == 1:
            paragraph.text = original_text.replace(token, lines[0])
            continue

        if original_text.strip() == token:
            paragraph.text = lines[0]
            try:
                paragraph.style = "List Bullet"
            except KeyError:
                pass
            cursor = paragraph
            for line in lines[1:]:
                cursor = _insert_paragraph_after(cursor)
                cursor.text = line
                try:
                    cursor.style = "List Bullet"
                except KeyError:
                    pass
            continue

        bullet_text = "\n".join(f"- {line}" for line in lines)
        paragraph.text = original_text.replace(token, bullet_text)

    return len(matched)


def query_to_text_payload(mart_rows: Sequence[Mapping[str, Any]]) -> list[str]:
    """Convert rows from a mart/view into short text lines.

    Adapter strategy:
    - if row has `text`, `summary` or `value` field: use it directly;
    - if row has `label` and `value`: use `label: value`;
    - fallback: first scalar fields as `field=value` fragments.

    Args:
        mart_rows: Sequence of rows returned by one mart/view query.

    Returns:
        List of short text lines ready for placeholder rendering.

    Raises:
        PlaceholderTextRenderError: If input is invalid or no text line can be derived.
    """

    if not isinstance(mart_rows, Sequence):
        raise PlaceholderTextRenderError(
            "mart_rows invalido: esperado sequencia de linhas do mart."
        )
    if not mart_rows:
        raise PlaceholderTextRenderError(
            "mart_rows invalido: lista vazia sem conteudo para texto."
        )

    lines: list[str] = []
    direct_keys = ("text", "summary", "value")
    for idx, row in enumerate(mart_rows):
        if not isinstance(row, Mapping):
            raise PlaceholderTextRenderError(
                f"mart_rows[{idx}] invalido: esperado objeto/Mapping."
            )

        direct_value = None
        for key in direct_keys:
            value = row.get(key)
            if isinstance(value, str) and value.strip():
                direct_value = value.strip()
                break
        if direct_value:
            lines.append(direct_value)
            continue

        label = row.get("label")
        value = row.get("value")
        if isinstance(label, str) and label.strip() and value is not None:
            lines.append(f"{label.strip()}: {value}")
            continue

        parts: list[str] = []
        for key, value in row.items():
            if isinstance(value, (str, int, float, bool)) and str(value).strip():
                parts.append(f"{key}={value}")
            if len(parts) == 3:
                break
        if parts:
            lines.append(", ".join(parts))

    clean_lines = [line for line in lines if line.strip()]
    if not clean_lines:
        raise PlaceholderTextRenderError(
            "Nao foi possivel derivar texto layout-ready a partir de mart_rows."
        )
    return clean_lines
