"""Run TMJ 2025 dry-run pipeline end-to-end using local fixtures.

This script executes a local, fixture-driven flow:
1. Run extractors (XLSX/PDF/PPTX) into staging artifacts.
2. Seed a local SQLite registry/session catalog.
3. Run DQ + show coverage preflight.
4. Render a DOCX report and emit final artifacts.

Expected outputs (in ``--out-dir``):
- ``dq_report.json``
- ``coverage_report.json``
- ``report.docx``
- ``manifest.json``
"""

from __future__ import annotations

import argparse
import shutil
import sqlite3
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
import zipfile

from docx import Document
from sqlmodel import Session, SQLModel, create_engine


REPO_ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = REPO_ROOT / "backend"
PROJECT_PARENT = REPO_ROOT.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))
if str(PROJECT_PARENT) not in sys.path:
    sys.path.insert(0, str(PROJECT_PARENT))
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.models.etl_registry import IngestionRun, IngestionStatus, Source, SourceKind
from etl.extract.extract_pdf_assisted import extract_pdf_assisted
from etl.extract.extract_pptx_slide_text import extract_pptx_slide_text
from etl.extract.extract_xlsx_leads_festival import extract_leads_xlsx
from etl.extract.extract_xlsx_optin_aceitos import extract_optin_xlsx
from etl.transform.agenda_loader import load_agenda_master
from etl.validate import cli_dq
from reports.word import cli_report


@dataclass(frozen=True)
class FixtureInputs:
    """Resolved fixture inputs for dry-run execution.

    Args:
        optin_xlsx: Fixture XLSX path for opt-in extraction.
        leads_xlsx: Fixture XLSX path for leads extraction.
        access_pdf: Fixture PDF path for assisted access-control extraction.
        agenda_master: Fixture agenda master path.
    """

    optin_xlsx: Path
    leads_xlsx: Path
    access_pdf: Path
    agenda_master: Path


@dataclass(frozen=True)
class ExtractedSource:
    """One extracted source that should be registered in ingestion catalog."""

    source_id: str
    kind: SourceKind
    source_path: Path
    extractor_name: str
    status: IngestionStatus


def _resolve_path(path_value: str | Path, *, base_dir: Path) -> Path:
    """Resolve a path string against repository base when relative.

    Args:
        path_value: Raw path from CLI args.
        base_dir: Base directory for relative paths.

    Returns:
        Absolute normalized path.
    """

    path = Path(path_value)
    if not path.is_absolute():
        path = base_dir / path
    return path.resolve()


def _require_fixture(path: Path, *, label: str) -> Path:
    """Validate fixture existence with actionable error.

    Args:
        path: Fixture file path.
        label: Human-readable fixture label.

    Returns:
        The same path when valid.

    Raises:
        FileNotFoundError: If fixture path does not exist.
    """

    if not path.exists():
        raise FileNotFoundError(
            f"Fixture nao encontrado ({label}): {path}. "
            "Como corrigir: informar caminho valido via argumento CLI."
        )
    return path


def _write_minimal_pptx_fixture(path: Path) -> Path:
    """Create a minimal PPTX (ZIP/XML) fixture used by dry-run extraction.

    Args:
        path: Output PPTX path.

    Returns:
        Created PPTX path.
    """

    slide_1_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:nvSpPr><p:nvPr><p:ph type="title"/></p:nvPr></p:nvSpPr>
        <p:txBody><a:p><a:r><a:t>Midias e Alcance</a:t></a:r></a:p></p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:nvPr/></p:nvSpPr>
        <p:txBody><a:p><a:r><a:t>Impressoes totais</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
"""
    slide_2_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:nvSpPr><p:nvPr><p:ph type="title"/></p:nvPr></p:nvSpPr>
        <p:txBody><a:p><a:r><a:t>Social Listening</a:t></a:r></a:p></p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:nvPr/></p:nvSpPr>
        <p:txBody><a:p><a:r><a:t>Sentimento positivo 78%</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
"""

    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("ppt/slides/slide1.xml", slide_1_xml)
        archive.writestr("ppt/slides/slide2.xml", slide_2_xml)
    return path


def _write_minimal_report_template(path: Path) -> Path:
    """Create a minimal DOCX template with required placeholders.

    Args:
        path: Target template path.

    Returns:
        Created template path.
    """

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("TMJ 2025 - Dry Run", level=1)
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.add_paragraph("{{SHOW__COVERAGE__TABLE}}")
    document.save(path)
    return path


def _create_event_sessions_table(engine) -> None:  # noqa: ANN001
    """Create minimal ``event_sessions`` table required by coverage evaluator.

    Args:
        engine: SQLAlchemy/SQLModel engine.
    """

    with engine.begin() as conn:
        conn.exec_driver_sql(
            """
            CREATE TABLE IF NOT EXISTS event_sessions (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              event_id INTEGER NULL,
              session_key TEXT NOT NULL,
              session_name TEXT NOT NULL,
              session_type TEXT NOT NULL,
              session_date TEXT NOT NULL,
              session_start_at TEXT NULL,
              session_end_at TEXT NULL,
              source_of_truth_source_id TEXT NULL,
              created_at TEXT NULL,
              updated_at TEXT NULL
            );
            """
        )


def _seed_event_sessions_from_agenda(
    *,
    engine,
    agenda_path: Path,
    event_id: int,
) -> None:
    """Insert expected sessions from master agenda into SQLite table.

    Args:
        engine: SQLAlchemy/SQLModel engine.
        agenda_path: Agenda master YAML/CSV path.
        event_id: Event identifier filter.

    Raises:
        ValueError: If agenda has no sessions for target event.
    """

    agenda = load_agenda_master(agenda_path)
    sessions = [session for session in agenda.sessions if int(session.event_id) == int(event_id)]
    if not sessions:
        raise ValueError(
            f"Agenda master sem sessoes para event_id={event_id}: {agenda_path}."
        )

    rows = []
    for session in sessions:
        session_type = str(session.session_type.value).strip().upper()
        if not session_type:
            session_type = "OUTRO"
        session_key = (
            str(session.session_key).strip()
            if session.session_key
            else f"TMJ2025_{session.session_date.strftime('%Y%m%d')}_{session_type}"
        )
        rows.append(
            {
                "event_id": int(session.event_id),
                "session_key": session_key,
                "session_name": str(session.name),
                "session_type": session_type,
                "session_date": session.session_date.isoformat(),
                "session_start_at": session.start_at.isoformat(),
            }
        )

    insert_sql = """
        INSERT INTO event_sessions (
          event_id,
          session_key,
          session_name,
          session_type,
          session_date,
          session_start_at
        ) VALUES (
          :event_id,
          :session_key,
          :session_name,
          :session_type,
          :session_date,
          :session_start_at
        );
    """
    with engine.begin() as conn:
        for row in rows:
            conn.exec_driver_sql(insert_sql, row)


def _seed_registry(
    *,
    engine,
    extracted_sources: Iterable[ExtractedSource],
) -> int:
    """Seed ``sources`` and ``ingestions`` tables for dry-run preflight.

    Args:
        engine: SQLModel engine where registry tables exist.
        extracted_sources: Sources extracted during dry-run.

    Returns:
        Ingestion ID to be used by preflight DQ command.
    """

    first_ingestion_id: int | None = None
    with Session(engine) as session:
        for item in extracted_sources:
            source = Source(
                source_id=item.source_id,
                kind=item.kind,
                uri=str(item.source_path),
            )
            session.add(source)
            session.commit()
            session.refresh(source)

            run = IngestionRun(
                source_pk=int(source.id),
                status=item.status,
                extractor_name=item.extractor_name,
                notes="dry-run fixture ingestion",
                records_read=1,
                records_loaded=1,
            )
            session.add(run)
            session.commit()
            session.refresh(run)
            if first_ingestion_id is None:
                first_ingestion_id = int(run.id)

    if first_ingestion_id is None:
        raise ValueError("Falha ao criar ingestao de dry-run.")
    return first_ingestion_id


def _run_extractors(
    *,
    fixtures: FixtureInputs,
    out_dir: Path,
    generated_pptx_path: Path,
) -> list[ExtractedSource]:
    """Run fixture extractors and return extracted source registry metadata.

    Args:
        fixtures: Fixture input paths.
        out_dir: Output staging directory.
        generated_pptx_path: Path where temporary PPTX fixture is generated.

    Returns:
        List of extracted source metadata for registry seed.
    """

    out_dir.mkdir(parents=True, exist_ok=True)
    _write_minimal_pptx_fixture(generated_pptx_path)

    extract_optin_xlsx(
        source_id="SRC_XLSX_OPTIN_NOTURNO_DOZE",
        xlsx_path=fixtures.optin_xlsx,
        out_dir=out_dir,
        include_pii=False,
    )
    extract_leads_xlsx(
        source_id="SRC_XLSX_LEADS_DIURNO_DOZE",
        xlsx_path=fixtures.leads_xlsx,
        out_dir=out_dir,
        include_pii=False,
    )
    extract_pptx_slide_text(
        source_id="SRC_PPTX_SOCIAL_NOTURNO_TREZE",
        pptx_path=generated_pptx_path,
        out_dir=out_dir,
    )
    extract_pdf_assisted(
        source_id="SRC_PDF_ACESSO_NOTURNO_TREZE",
        pdf_path=fixtures.access_pdf,
        out_dir=out_dir,
        template="access_control",
    )

    return [
        ExtractedSource(
            source_id="SRC_XLSX_OPTIN_NOTURNO_DOZE",
            kind=SourceKind.XLSX,
            source_path=fixtures.optin_xlsx,
            extractor_name="extract_xlsx_optin_aceitos",
            status=IngestionStatus.SUCCESS,
        ),
        ExtractedSource(
            source_id="SRC_XLSX_LEADS_DIURNO_DOZE",
            kind=SourceKind.XLSX,
            source_path=fixtures.leads_xlsx,
            extractor_name="extract_xlsx_leads_festival",
            status=IngestionStatus.SUCCESS,
        ),
        ExtractedSource(
            source_id="SRC_PPTX_SOCIAL_NOTURNO_TREZE",
            kind=SourceKind.PPTX,
            source_path=generated_pptx_path,
            extractor_name="extract_pptx_slide_text",
            status=IngestionStatus.SUCCESS,
        ),
        ExtractedSource(
            source_id="SRC_PDF_ACESSO_NOTURNO_TREZE",
            kind=SourceKind.PDF,
            source_path=fixtures.access_pdf,
            extractor_name="extract_pdf_assisted",
            status=IngestionStatus.SUCCESS,
        ),
    ]


def _build_fixture_inputs(args: argparse.Namespace) -> FixtureInputs:
    """Resolve and validate dry-run fixture inputs from CLI args.

    Args:
        args: Parsed CLI namespace.

    Returns:
        Validated fixture input paths.

    Raises:
        FileNotFoundError: If any required fixture path is missing.
    """

    fixtures_root = _resolve_path(args.fixtures_root, base_dir=REPO_ROOT)
    optin = _resolve_path(args.optin_xlsx, base_dir=REPO_ROOT)
    leads = _resolve_path(args.leads_xlsx, base_dir=REPO_ROOT)
    pdf = _resolve_path(args.access_pdf, base_dir=REPO_ROOT)
    agenda = _resolve_path(args.agenda_path, base_dir=REPO_ROOT)

    # Resolve defaults relative to fixtures root only when user kept defaults.
    if str(args.optin_xlsx) == "tests/fixtures/xlsx/optin_min.xlsx":
        optin = fixtures_root / "xlsx" / "optin_min.xlsx"
    if str(args.leads_xlsx) == "tests/fixtures/xlsx/leads_min.xlsx":
        leads = fixtures_root / "xlsx" / "leads_min.xlsx"
    if str(args.access_pdf) == "tests/fixtures/pdf/min_table.pdf":
        pdf = fixtures_root / "pdf" / "min_table.pdf"
    if str(args.agenda_path) == "tests/fixtures/agenda/agenda_master_min.yml":
        agenda = fixtures_root / "agenda" / "agenda_master_min.yml"

    return FixtureInputs(
        optin_xlsx=_require_fixture(optin, label="optin_xlsx"),
        leads_xlsx=_require_fixture(leads, label="leads_xlsx"),
        access_pdf=_require_fixture(pdf, label="access_pdf"),
        agenda_master=_require_fixture(agenda, label="agenda_master"),
    )


def _build_parser() -> argparse.ArgumentParser:
    """Build CLI parser for TMJ dry-run script."""

    parser = argparse.ArgumentParser(prog="dry_run_tmj_pipeline")
    parser.add_argument(
        "--out-dir",
        default="reports/tmj2025/dry_run",
        help="Output directory for artifacts and temporary files.",
    )
    parser.add_argument(
        "--event-id",
        type=int,
        default=2025,
        help="Event identifier used in preflight/report generation.",
    )
    parser.add_argument(
        "--fixtures-root",
        default="tests/fixtures",
        help="Root directory containing fixture files.",
    )
    parser.add_argument(
        "--optin-xlsx",
        default="tests/fixtures/xlsx/optin_min.xlsx",
        help="Fixture XLSX path for opt-in extractor.",
    )
    parser.add_argument(
        "--leads-xlsx",
        default="tests/fixtures/xlsx/leads_min.xlsx",
        help="Fixture XLSX path for leads extractor.",
    )
    parser.add_argument(
        "--access-pdf",
        default="tests/fixtures/pdf/min_table.pdf",
        help="Fixture PDF path for assisted access-control extractor.",
    )
    parser.add_argument(
        "--agenda-path",
        default="tests/fixtures/agenda/agenda_master_min.yml",
        help="Agenda master fixture path (YAML/CSV).",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    """Run dry-run ETL -> DQ -> coverage -> report flow.

    Args:
        argv: Optional CLI argument list.

    Returns:
        Exit code (``0`` on success, ``1`` on expected execution errors).

    Raises:
        RuntimeError: If report render command fails during dry-run.
    """

    parser = _build_parser()
    args = parser.parse_args(argv)

    try:
        out_dir = _resolve_path(args.out_dir, base_dir=REPO_ROOT)
        out_dir.mkdir(parents=True, exist_ok=True)

        fixtures = _build_fixture_inputs(args)

        staging_dir = out_dir / "staging"
        generated_pptx = out_dir / "fixtures_runtime" / "social_min.pptx"
        template_path = out_dir / "template_dry_run.docx"
        output_report_path = out_dir / "report.docx"
        dq_report_path = out_dir / "dq_report.json"
        coverage_report_path = out_dir / "coverage_report.json"

        db_path = out_dir / "dry_run.sqlite"
        if db_path.exists():
            db_path.unlink()
        sqlite_url = f"sqlite:///{db_path.as_posix()}"
        engine = create_engine(sqlite_url, connect_args={"check_same_thread": False})

        extracted_sources = _run_extractors(
            fixtures=fixtures,
            out_dir=staging_dir,
            generated_pptx_path=generated_pptx,
        )

        SQLModel.metadata.create_all(
            engine,
            tables=[Source.__table__, IngestionRun.__table__],
        )
        _create_event_sessions_table(engine)
        _seed_event_sessions_from_agenda(
            engine=engine,
            agenda_path=fixtures.agenda_master,
            event_id=int(args.event_id),
        )
        ingestion_id = _seed_registry(engine=engine, extracted_sources=extracted_sources)

        _write_minimal_report_template(template_path)

        # Force DQ preflight to use local dry-run SQLite, not external DB.
        cli_dq.engine = engine

        rc = cli_report.main(
            [
                "report:render",
                "--event-id",
                str(int(args.event_id)),
                "--template-path",
                str(template_path),
                "--output-path",
                str(output_report_path),
                "--run-preflight",
                "--ingestion-id",
                str(ingestion_id),
                "--dq-out-json",
                str(dq_report_path),
                "--coverage-out-json",
                str(coverage_report_path),
                "--agenda-master-path",
                str(fixtures.agenda_master),
            ]
        )
        if int(rc) != 0:
            raise RuntimeError(
                "Falha no report:render durante dry-run. "
                "Como corrigir: revisar logs de preflight DQ/coverage."
            )

        report_manifest_path = out_dir / "report_manifest.json"
        manifest_path = out_dir / "manifest.json"
        if not report_manifest_path.exists():
            raise FileNotFoundError(
                f"Manifest do renderer nao encontrado: {report_manifest_path}"
            )
        shutil.copyfile(report_manifest_path, manifest_path)

        required_outputs = (
            dq_report_path,
            coverage_report_path,
            output_report_path,
            manifest_path,
        )
        missing = [path for path in required_outputs if not path.exists()]
        if missing:
            missing_list = ", ".join(str(path) for path in missing)
            raise RuntimeError(
                "Dry-run concluiu sem todos os artefatos esperados. "
                f"Ausentes: {missing_list}"
            )

        print("[OK] dry-run TMJ pipeline concluido")
        print(f" - dq_report: {dq_report_path}")
        print(f" - coverage_report: {coverage_report_path}")
        print(f" - report_docx: {output_report_path}")
        print(f" - manifest: {manifest_path}")
        return 0
    except (FileNotFoundError, ValueError, RuntimeError, OSError, sqlite3.Error) as exc:
        print(f"[ERROR] {exc}")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
