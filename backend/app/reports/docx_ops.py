"""Low-level helpers for editing DOCX templates with python-docx.

python-docx has limited APIs for:
- deleting blocks (paragraphs/tables),
- inserting tables at arbitrary positions.

This module provides small, focused helpers for those operations.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Optional

from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
import unicodedata
import re


@dataclass(frozen=True)
class HeadingMatch:
    paragraph: Paragraph
    level: int
    text: str


def heading_level(paragraph: Paragraph) -> Optional[int]:
    name = getattr(paragraph.style, "name", "") or ""
    if not name.startswith("Heading"):
        # Fallback: some templates (or localized Word installs) may not have
        # "Heading N" style names. When that happens, infer heading levels from
        # the leading numbering token: "4." -> level 1, "4.1" -> level 2, etc.
        text = (paragraph.text or "").strip()
        if not text:
            return None
        token = text.split()[0]
        token_raw = token
        token = token.rstrip(".")
        if "." not in token_raw and not token_raw.endswith("."):
            return None
        if not re.match(r"^\\d+(?:\\.\\d+)*$", token):
            return None
        parts = [p for p in token.split(".") if p]
        return len(parts) if parts else None
    parts = name.split()
    if len(parts) != 2:
        return None
    try:
        return int(parts[1])
    except ValueError:
        return None


def find_heading_by_prefix(doc: DocxDocument, *, prefix: str) -> HeadingMatch:
    """Find the first heading paragraph whose text starts with `prefix`."""
    pref = (prefix or "").strip()
    if not pref:
        raise ValueError("prefix vazio")

    def norm(s: str) -> str:
        s = s or ""
        n = unicodedata.normalize("NFKD", s)
        a = "".join(ch for ch in n if ord(ch) < 128)
        return a.strip().lower()

    pref_norm = norm(pref)
    for p in doc.paragraphs:
        lvl = heading_level(p)
        if lvl is None:
            continue
        text = (p.text or "").strip()
        if norm(text).startswith(pref_norm):
            return HeadingMatch(paragraph=p, level=lvl, text=text)

    raise ValueError(f"Heading nao encontrado: prefix={pref!r}")


def find_next_heading(
    doc: DocxDocument,
    *,
    after: Paragraph,
    max_level: int,
) -> Optional[HeadingMatch]:
    """Find the next heading after a given paragraph with level <= max_level."""
    if max_level <= 0:
        raise ValueError("max_level invalido")

    seen = False
    for p in doc.paragraphs:
        if not seen:
            # `doc.paragraphs` creates new wrapper objects on each access; use the
            # underlying XML element identity instead of Python object identity.
            if p._element is after._element:
                seen = True
            continue
        lvl = heading_level(p)
        if lvl is None:
            continue
        if lvl <= max_level:
            text = (p.text or "").strip()
            return HeadingMatch(paragraph=p, level=lvl, text=text)
    return None


def _body_children(doc: DocxDocument):
    body = doc.element.body
    return body, list(body.iterchildren())


def clear_blocks_between(
    doc: DocxDocument,
    *,
    start: Paragraph,
    end: Paragraph | None,
) -> None:
    """Remove all body blocks between start paragraph (exclusive) and end paragraph (exclusive).

    This removes both paragraphs and tables.
    """
    body, children = _body_children(doc)
    start_el = start._element
    try:
        start_idx = children.index(start_el)
    except ValueError as exc:
        raise ValueError("start paragraph element nao encontrado no body") from exc

    if end is None:
        end_idx = len(children)
        # Preserve the final section properties element, if present.
        if children and children[-1].tag.endswith("}sectPr"):
            end_idx = len(children) - 1
    else:
        end_el = end._element
        try:
            end_idx = children.index(end_el)
        except ValueError as exc:
            raise ValueError("end paragraph element nao encontrado no body") from exc

    # Remove in forward order from the snapshot list.
    for el in children[start_idx + 1 : end_idx]:
        body.remove(el)


def insert_paragraph_after(
    doc: DocxDocument,
    *,
    anchor: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    """Insert a new paragraph after `anchor` and return it."""
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    anchor._element.addnext(p._element)
    return p


def insert_paragraph_after_block(
    doc: DocxDocument,
    *,
    anchor: Paragraph | Table,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    """Insert a paragraph after a paragraph OR a table."""
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    anchor._element.addnext(p._element)  # type: ignore[attr-defined]
    return p


def insert_table_after(
    doc: DocxDocument,
    *,
    anchor: Paragraph,
    rows: int,
    cols: int,
    style: str | None = "Table Grid",
) -> Table:
    """Insert an empty table after `anchor` and return it."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if style:
        tbl.style = style
    anchor._element.addnext(tbl._element)
    return tbl


def insert_table_after_block(
    doc: DocxDocument,
    *,
    anchor: Paragraph | Table,
    rows: int,
    cols: int,
    style: str | None = "Table Grid",
) -> Table:
    """Insert a table after a paragraph OR a table."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if style:
        tbl.style = style
    anchor._element.addnext(tbl._element)  # type: ignore[attr-defined]
    return tbl


def set_table_header_bold(tbl: Table) -> None:
    if not tbl.rows:
        return
    for cell in tbl.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_kv_table(
    doc: DocxDocument,
    *,
    anchor: Paragraph | Table,
    rows: list[tuple[str, str]],
    style: str | None = "Table Grid",
) -> Table:
    """Insert a 2-col table after anchor with (label, value) rows."""
    tbl = insert_table_after_block(doc, anchor=anchor, rows=1, cols=2, style=style)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metrica"
    hdr[1].text = "Valor"
    set_table_header_bold(tbl)
    for label, value in rows:
        r = tbl.add_row().cells
        r[0].text = str(label)
        r[1].text = str(value)
    return tbl


def add_matrix_table(
    doc: DocxDocument,
    *,
    anchor: Paragraph | Table,
    header: list[str],
    rows: list[list[str]],
    style: str | None = "Table Grid",
) -> Table:
    """Insert a matrix table (header row + N rows) after anchor."""
    tbl = insert_table_after_block(doc, anchor=anchor, rows=1, cols=len(header), style=style)
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = str(h)
    set_table_header_bold(tbl)
    for row in rows:
        r = tbl.add_row().cells
        for i, v in enumerate(row):
            r[i].text = str(v)
    return tbl
