"""Unit tests for Word template placeholder validation helpers."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from reports.word.placeholders_mapping import load_placeholders_mapping
from reports.word.template_validate import (
    find_placeholders,
    validate_template_placeholders,
)


def test_find_placeholders_extracts_tokens_from_paragraphs_and_tables(tmp_path: Path) -> None:
    """Finder should extract unique placeholder ids from body and table cells."""

    template_path = tmp_path / "template_find.docx"
    document = Document()
    document.add_paragraph("Resumo: {{FONTES__SUMMARY__TEXT}}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{PUBLICO__ATTENDANCE__TABLE}}"
    document.save(template_path)

    placeholders = find_placeholders(template_path)

    assert placeholders == {"FONTES__SUMMARY__TEXT", "PUBLICO__ATTENDANCE__TABLE"}


def test_validate_template_placeholders_returns_unknown_findings(tmp_path: Path) -> None:
    """Validator should report template placeholders that are not mapped."""

    mapping = load_placeholders_mapping()
    placeholders = {"FONTES__SUMMARY__TEXT", "DESCONHECIDO__BLOCO__TEXT"}

    findings = validate_template_placeholders(placeholders, mapping)

    assert len(findings) == 1
    assert findings[0].code == "UNKNOWN_PLACEHOLDER"
    assert findings[0].placeholder_id == "DESCONHECIDO__BLOCO__TEXT"
