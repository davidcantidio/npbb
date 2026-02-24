"""Tests for table/figure renderers and chart PNG generation."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from reports.charts import ChartSeries, ChartSpec, ChartType, render_chart_png
from reports.word import (
    PlaceholderFigureRenderError,
    PlaceholderTableRenderError,
    render_figure_placeholder,
    render_table_placeholder,
)


def _sample_chart_rows() -> list[dict[str, object]]:
    """Return deterministic sample rows for chart and figure tests."""

    return [
        {"dia": "2025-12-12", "presentes": 120},
        {"dia": "2025-12-13", "presentes": 180},
        {"dia": "2025-12-14", "presentes": 150},
    ]


def test_render_table_placeholder_respects_contract_columns(tmp_path: Path) -> None:
    """Table renderer should create DOCX table with expected contract columns."""

    placeholder_id = "PUBLICO__ATTENDANCE__TABLE"
    template_path = tmp_path / "table_template.docx"
    document = Document()
    document.add_paragraph(f"{{{{{placeholder_id}}}}}")
    document.save(template_path)

    doc = Document(template_path)
    replaced = render_table_placeholder(
        doc,
        placeholder_id,
        columns=["sessao", "presentes", "validos"],
        rows=[
            {"sessao": "show_12", "presentes": 100, "validos": 120},
            {"sessao": "show_13", "presentes": 150, "validos": 180},
        ],
    )
    out_path = tmp_path / "table_out.docx"
    doc.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    assert len(rendered.tables) == 1
    header = [cell.text for cell in rendered.tables[0].rows[0].cells]
    assert header == ["sessao", "presentes", "validos"]


def test_render_table_placeholder_fails_when_contract_column_missing(tmp_path: Path) -> None:
    """Table renderer should fail when one required contract column is missing."""

    placeholder_id = "PUBLICO__ATTENDANCE__TABLE"
    template_path = tmp_path / "table_contract_fail.docx"
    document = Document()
    document.add_paragraph(f"{{{{{placeholder_id}}}}}")
    document.save(template_path)

    doc = Document(template_path)
    with pytest.raises(PlaceholderTableRenderError, match="Contrato de colunas violado"):
        render_table_placeholder(
            doc,
            placeholder_id,
            columns=["sessao", "presentes", "validos"],
            rows=[{"sessao": "show_12", "presentes": 100}],
        )


def test_render_chart_png_and_figure_placeholder_insert_image(tmp_path: Path) -> None:
    """Chart generator + figure renderer should create and insert PNG in DOCX."""

    chart_spec = ChartSpec(
        chart_type=ChartType.BAR,
        title="Presencas por dia",
        x_label="Dia",
        y_label="Presentes",
        series=(ChartSeries(name="Presencas", x_field="dia", y_field="presentes"),),
    )

    png_a = render_chart_png(chart_spec, _sample_chart_rows())
    png_b = render_chart_png(chart_spec, _sample_chart_rows())
    assert isinstance(png_a, bytes)
    assert png_a[:8] == b"\x89PNG\r\n\x1a\n"
    assert png_a == png_b

    image_path = tmp_path / "figure.png"
    image_path.write_bytes(png_a)

    placeholder_id = "PRESALE__CURVES__FIGURE"
    template_path = tmp_path / "figure_template.docx"
    template = Document()
    template.add_paragraph(f"{{{{{placeholder_id}}}}}")
    template.save(template_path)

    doc = Document(template_path)
    replaced = render_figure_placeholder(
        doc,
        placeholder_id,
        image_path,
        "Figura 1 - Presencas por dia",
    )
    out_path = tmp_path / "figure_out.docx"
    doc.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert f"{{{{{placeholder_id}}}}}" not in text
    assert "Figura 1 - Presencas por dia" in text
    assert len(rendered.inline_shapes) == 1


def test_render_figure_placeholder_fails_for_missing_image(tmp_path: Path) -> None:
    """Figure renderer should fail with actionable error when image is missing."""

    doc = Document()
    doc.add_paragraph("{{PRESALE__CURVES__FIGURE}}")

    with pytest.raises(PlaceholderFigureRenderError, match="nao encontrada"):
        render_figure_placeholder(
            doc,
            "PRESALE__CURVES__FIGURE",
            tmp_path / "missing.png",
            "Figura X",
        )
