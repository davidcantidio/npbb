"""Figure placeholder rendering helpers for Word report sections.

This module provides figure rendering utilities to replace placeholders
inside a DOCX document with image + caption output.
"""

from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from .render_text import placeholder_token


class PlaceholderFigureRenderError(ValueError):
    """Raised when figure placeholder rendering fails."""


def _iter_document_paragraphs(doc: DocxDocument) -> list[Paragraph]:
    """Collect all paragraphs from document body and table cells.

    Args:
        doc: Open DOCX document.

    Returns:
        Flat list with body and table-cell paragraphs.
    """

    paragraphs: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    """Insert a paragraph right after an existing paragraph.

    Args:
        anchor: Paragraph that will receive a new sibling after it.

    Returns:
        Newly created paragraph object.
    """

    new_node = OxmlElement("w:p")
    anchor._p.addnext(new_node)
    return Paragraph(new_node, anchor._parent)


def render_figure_placeholder(
    doc: DocxDocument,
    placeholder_id: str,
    image_path: Path | str,
    caption: str,
    *,
    width_inches: float = 6.0,
) -> int:
    """Render one figure placeholder in a DOCX document.

    Rules:
    - placeholder token format is `{{PLACEHOLDER_ID}}`;
    - each occurrence becomes one image block + one caption paragraph;
    - image width is fixed by `width_inches` for layout consistency.

    Args:
        doc: Open DOCX document.
        placeholder_id: Placeholder identifier to replace.
        image_path: PNG/JPG image path.
        caption: Figure caption text.
        width_inches: Output image width in inches.

    Returns:
        Number of placeholder occurrences replaced.

    Raises:
        PlaceholderFigureRenderError: If placeholder/image/caption is invalid.
    """

    token = placeholder_token(placeholder_id)
    source_path = Path(image_path)
    if not source_path.exists():
        raise PlaceholderFigureRenderError(f"Imagem da figura nao encontrada: {source_path}")
    if not str(caption or "").strip():
        raise PlaceholderFigureRenderError(
            f"Caption invalida para figura {placeholder_id}: texto vazio."
        )
    if float(width_inches) <= 0:
        raise PlaceholderFigureRenderError(
            f"width_inches invalido para figura {placeholder_id}: esperado valor positivo."
        )

    paragraphs = _iter_document_paragraphs(doc)
    matched = [paragraph for paragraph in paragraphs if token in paragraph.text]
    if not matched:
        raise PlaceholderFigureRenderError(
            f"Placeholder de figura nao encontrado no DOCX: {placeholder_id}."
        )

    for paragraph in matched:
        paragraph.text = paragraph.text.replace(token, "").strip()

        image_paragraph = _insert_paragraph_after(paragraph)
        image_run = image_paragraph.add_run()
        image_run.add_picture(str(source_path), width=Inches(width_inches))

        caption_paragraph = _insert_paragraph_after(image_paragraph)
        caption_paragraph.text = caption.strip()

    return len(matched)
