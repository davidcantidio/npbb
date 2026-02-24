"""Unit tests for requirements-to-schema mapping markdown renderer."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from core.spec.mapping_models import parse_mapping_contract
from core.spec.render_mapping_md import render_mapping_md
from core.spec.spec_models import DocxChecklistRow
from etl.cli_spec import main


def _checklist_fixture() -> list[DocxChecklistRow]:
    """Return checklist fixture with public-gauge relevant items."""
    return [
        DocxChecklistRow(
            section="Publico do evento (controle de acesso - entradas validadas)",
            expected_content="Entradas validadas",
            granularity="evento_dia_sessao",
            required_fields="attendance_access_control.entries_validated",
            source_in_docx='Secao "Publico"',
            status="GAP",
        ),
        DocxChecklistRow(
            section="Dinamica de vendas (pre-venda) - shows (Opt-in aceitos)",
            expected_content="Opt-in aceitos",
            granularity="evento_dia_sessao",
            required_fields="optin_transactions.optin_accepted",
            source_in_docx='Secao "Pre-venda"',
            status="GAP",
        ),
        DocxChecklistRow(
            section="Publico unico deduplicado",
            expected_content="Publico unico",
            granularity="evento_dia_sessao",
            required_fields="audience_metrics.unique_public",
            source_in_docx='Secao "Publico"',
            status="GAP",
        ),
    ]


def _mapping_fixture():
    """Return typed mapping fixture equivalent to YAML content."""
    return parse_mapping_contract(
        {
            "schema_version": "1.0",
            "spec_source": "docs/analises/eventos/tamo_junto_2025/planning/00_docx_as_spec.md",
            "requirements": [
                {
                    "requirement_id": "REQ-ENT-001",
                    "item_docx": "Publico do evento (controle de acesso - entradas validadas)",
                    "target_field": "attendance_access_control.entries_validated",
                    "calculation_rule": "SUM(entries_validated) por sessao",
                    "sources": [
                        {"source_id": "pdf_access_control_12_12", "location": "pagina 1 / tabela Presencas"}
                    ],
                    "validations": [
                        {
                            "rule_id": "VAL-ENT-001",
                            "rule_type": "reconciliation",
                            "severity": "error",
                            "expression": "entries_validated >= 0",
                        }
                    ],
                },
                {
                    "requirement_id": "REQ-OPT-001",
                    "item_docx": "Dinamica de vendas (pre-venda) - shows (Opt-in aceitos)",
                    "target_field": "optin_transactions.optin_accepted",
                    "calculation_rule": "SUM(optin_accepted) por sessao",
                    "sources": [
                        {"source_id": "xlsx_optin_aceitos_2025_12", "location": "aba Vendas / colunas A:D"}
                    ],
                    "validations": [
                        {
                            "rule_id": "VAL-OPT-001",
                            "rule_type": "not_null",
                            "severity": "error",
                            "expression": "optin_accepted IS NOT NULL",
                        }
                    ],
                },
            ],
            "marts": [
                {
                    "mart_name": "mart_report_show_day_summary",
                    "grain": "evento_dia_sessao",
                    "description": "Resumo por sessao",
                    "requirement_ids": ["REQ-ENT-001", "REQ-OPT-001"],
                    "output_fields": [
                        {
                            "field": "entries_validated",
                            "source": "public.attendance_access_control.entries_validated",
                            "rule": "Soma por sessao",
                        },
                        {
                            "field": "optin_accepted",
                            "source": "public.optin_transactions.optin_accepted",
                            "rule": "Soma por sessao",
                        },
                    ],
                }
            ],
        }
    )


def test_render_mapping_md_reflects_mapping_yaml_and_public_gauge_notes() -> None:
    """Rendered markdown contains mapping data and public-gauge notes."""
    markdown = render_mapping_md(_checklist_fixture(), _mapping_fixture())

    assert "REQ-ENT-001" in markdown
    assert "public.attendance_access_control.entries_validated" in markdown
    assert "VAL-ENT-001 [reconciliation/error] entries_validated >= 0" in markdown
    assert "pdf_access_control_12_12 (pagina 1 / tabela Presencas)" in markdown
    assert "Regua de publico: entradas validadas" in markdown
    assert "Regua de publico: opt-in aceitos" in markdown
    assert "| mart_report_show_day_summary | evento_dia_sessao | REQ-ENT-001, REQ-OPT-001 | entries_validated, optin_accepted |" in markdown


def test_render_mapping_md_marks_missing_checklist_items_as_gap() -> None:
    """Checklist section reports GAP when an item has no mapping."""
    markdown = render_mapping_md(_checklist_fixture(), _mapping_fixture())

    assert "| Publico unico deduplicado | GAP |" in markdown
    assert "| Publico do evento (controle de acesso - entradas validadas) | OK |" in markdown


def test_cli_spec_render_mapping_generates_output_file(tmp_path: Path) -> None:
    """CLI command `spec:render-mapping` writes mapping markdown output."""
    docx_path = tmp_path / "template.docx"
    mapping_path = tmp_path / "mapping.yml"
    output_path = tmp_path / "03_requirements_to_schema_mapping.md"

    doc = Document()
    doc.add_heading("Publico do evento (controle de acesso - entradas validadas)", level=1)
    doc.add_paragraph("Texto base.")
    doc.save(docx_path)

    mapping_path.write_text(
        """
schema_version: "1.0"
spec_source: "docs/analises/eventos/tamo_junto_2025/planning/00_docx_as_spec.md"
requirements:
  - requirement_id: "REQ-001"
    item_docx: "Publico do evento (controle de acesso - entradas validadas)"
    target_field: "attendance_access_control.entries_validated"
    calculation_rule: "SUM(entries_validated)"
    sources:
      - source_id: "pdf_access_control_12_12"
        location: "pagina 1 / tabela Presencas"
    validations:
      - rule_id: "VAL-001"
        rule_type: "reconciliation"
        severity: "error"
        expression: "entries_validated >= 0"
marts:
  - mart_name: "mart_report_show_day_summary"
    grain: "evento_dia_sessao"
    description: "Resumo por sessao"
    requirement_ids: ["REQ-001"]
    output_fields:
      - field: "entries_validated"
        source: "public.attendance_access_control.entries_validated"
        rule: "Soma por sessao"
""".strip(),
        encoding="utf-8",
    )

    rc = main(
        [
            "spec:render-mapping",
            "--docx",
            str(docx_path),
            "--mapping",
            str(mapping_path),
            "--out",
            str(output_path),
        ]
    )

    assert rc == 0
    assert output_path.exists()
    content = output_path.read_text(encoding="utf-8")
    assert "## Mapeamento Item -> Campo -> Regra" in content
    assert "public.attendance_access_control.entries_validated" in content
