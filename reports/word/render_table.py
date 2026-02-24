"""Table placeholder rendering helpers for Word report sections.

This module provides table rendering utilities to replace placeholders
inside a DOCX document with deterministic tabular output.
"""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from typing import Any

from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from .render_text import placeholder_token


class PlaceholderTableRenderError(ValueError):
    """Raised when table placeholder rendering fails."""


def _iter_document_paragraphs(doc: DocxDocument) -> list[Paragraph]:
    """Collect all paragraphs from document body and table cells.

    Args:
        doc: Open DOCX document.

    Returns:
        Flat list with body and table-cell paragraphs.
    """

    paragraphs: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _normalize_columns(columns: Sequence[str], *, placeholder_id: str) -> list[str]:
    """Normalize expected table columns for one placeholder.

    Args:
        columns: Expected columns in output order.
        placeholder_id: Placeholder identifier for error context.

    Returns:
        Normalized column names.

    Raises:
        PlaceholderTableRenderError: If column contract is invalid.
    """

    if not isinstance(columns, Sequence) or isinstance(columns, (str, bytes, bytearray)):
        raise PlaceholderTableRenderError(
            f"Payload de tabela invalido para {placeholder_id}: columns deve ser list[str]."
        )

    normalized = []
    for index, column in enumerate(columns):
        if not isinstance(column, str) or not column.strip():
            raise PlaceholderTableRenderError(
                f"Payload de tabela invalido para {placeholder_id}: "
                f"columns[{index}] deve ser texto nao vazio."
            )
        normalized.append(column.strip())

    if not normalized:
        raise PlaceholderTableRenderError(
            f"Payload de tabela invalido para {placeholder_id}: columns nao pode ser vazio."
        )

    duplicates = sorted({col for col in normalized if normalized.count(col) > 1})
    if duplicates:
        duplicate_list = ", ".join(duplicates)
        raise PlaceholderTableRenderError(
            f"Payload de tabela invalido para {placeholder_id}: colunas duplicadas ({duplicate_list})."
        )
    return normalized


def _normalize_rows(
    rows: Sequence[Mapping[str, Any] | Sequence[Any]],
    *,
    columns: Sequence[str],
    placeholder_id: str,
) -> list[list[str]]:
    """Normalize table rows and validate mart contract coverage.

    Args:
        rows: Row objects from mart payload.
        columns: Required columns (contract order).
        placeholder_id: Placeholder identifier for error context.

    Returns:
        Normalized rows as list of string cells.

    Raises:
        PlaceholderTableRenderError: If row format is invalid or contract is violated.
    """

    if not isinstance(rows, Sequence) or isinstance(rows, (str, bytes, bytearray)):
        raise PlaceholderTableRenderError(
            f"Payload de tabela invalido para {placeholder_id}: rows deve ser lista."
        )

    normalized_rows: list[list[str]] = []
    for row_index, row in enumerate(rows):
        if isinstance(row, Mapping):
            missing_columns = [column for column in columns if column not in row]
            if missing_columns:
                missing_list = ", ".join(missing_columns)
                raise PlaceholderTableRenderError(
                    f"Contrato de colunas violado para {placeholder_id}: "
                    f"rows[{row_index}] sem coluna(s) obrigatoria(s): {missing_list}."
                )
            values = [row.get(column) for column in columns]
        elif isinstance(row, Sequence) and not isinstance(row, (str, bytes, bytearray)):
            if len(row) != len(columns):
                raise PlaceholderTableRenderError(
                    f"Contrato de colunas violado para {placeholder_id}: "
                    f"rows[{row_index}] tem {len(row)} valores, esperado {len(columns)}."
                )
            values = list(row)
        else:
            raise PlaceholderTableRenderError(
                f"Payload de tabela invalido para {placeholder_id}: "
                f"rows[{row_index}] deve ser Mapping ou Sequence."
            )
        normalized_rows.append(["" if value is None else str(value) for value in values])
    return normalized_rows


def _render_table_after_paragraph(
    paragraph: Paragraph,
    *,
    columns: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> None:
    """Insert a table right after an anchor paragraph.

    Args:
        paragraph: Paragraph where table placeholder was found.
        columns: Header columns in output order.
        rows: Normalized row values.
    """

    try:
        table = paragraph._parent.add_table(
            rows=1 + len(rows),
            cols=len(columns),
            width=Inches(6.5),
        )
    except TypeError:
        table = paragraph._parent.add_table(rows=1 + len(rows), cols=len(columns))
    for col_index, column in enumerate(columns):
        table.cell(0, col_index).text = str(column)
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    paragraph._p.addnext(table._tbl)


def render_table_placeholder(
    doc: DocxDocument,
    placeholder_id: str,
    columns: Sequence[str],
    rows: Sequence[Mapping[str, Any] | Sequence[Any]],
) -> int:
    """Render one table placeholder in a DOCX document.

    Rules:
    - placeholder token format is `{{PLACEHOLDER_ID}}`;
    - table header follows expected `columns` order;
    - each row must satisfy column contract (all required columns present).

    Args:
        doc: Open DOCX document.
        placeholder_id: Placeholder identifier to replace.
        columns: Expected output columns in order.
        rows: Row payload (list of mapping or ordered list values).

    Returns:
        Number of placeholder occurrences replaced by tables.

    Raises:
        PlaceholderTableRenderError: If placeholder is missing or payload invalid.
    """

    token = placeholder_token(placeholder_id)
    normalized_columns = _normalize_columns(columns, placeholder_id=placeholder_id)
    normalized_rows = _normalize_rows(
        rows,
        columns=normalized_columns,
        placeholder_id=placeholder_id,
    )

    paragraphs = _iter_document_paragraphs(doc)
    matched = [paragraph for paragraph in paragraphs if token in paragraph.text]
    if not matched:
        raise PlaceholderTableRenderError(
            f"Placeholder de tabela nao encontrado no DOCX: {placeholder_id}."
        )

    for paragraph in matched:
        paragraph.text = paragraph.text.replace(token, "").strip()
        _render_table_after_paragraph(
            paragraph,
            columns=normalized_columns,
            rows=normalized_rows,
        )

    return len(matched)
