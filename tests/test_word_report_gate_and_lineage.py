"""Unit tests for report gap rendering, gate decisions, and manifest lineage."""

from __future__ import annotations

import json
from pathlib import Path
from textwrap import dedent

from docx import Document

from reports.word import (
    WordReportRenderer,
    evaluate_report_gate,
    load_dq_report,
    load_report_gate_policy,
    render_gaps_section,
)


def _fixture_path(*parts: str) -> Path:
    """Build absolute path under tests/fixtures."""

    return Path(__file__).resolve().parent / "fixtures" / Path(*parts)


def _write_policy(path: Path, content: str) -> Path:
    """Write temporary gate policy YAML fixture."""

    path.write_text(dedent(content).strip() + "\n", encoding="utf-8")
    return path


def test_render_gaps_section_includes_items_and_evidence_from_fixture() -> None:
    """GAP section should include findings, source, location, and evidence fields."""

    dq_report = load_dq_report(_fixture_path("dq", "dq_report_gaps_inconsistencies.json"))

    lines = render_gaps_section(dq_report)
    text = "\n".join(lines)

    assert "Resumo GAP/INCONSISTENTE" in text
    assert "INCONSISTENTE | dataset=show_12_dimac" in text
    assert "GAP | dataset=show_14_access_control" in text
    assert "source_id=SRC_DIMAC_2025" in text
    assert "source_id=SRC_ACCESS_2025" in text
    assert "location=page=7" in text
    assert "location=page=0" in text
    assert "Cobertura de shows por dia (12/12, 13/12, 14/12):" in text
    assert "12/12: INCONSISTENTE" in text
    assert "14/12: GAP" in text


def test_evaluate_report_gate_blocks_for_critical_show_gap(tmp_path: Path) -> None:
    """Gate should block output when a configured critical show gap is detected."""

    dq_report = load_dq_report(_fixture_path("dq", "dq_report_gaps_inconsistencies.json"))
    policy_path = _write_policy(
        tmp_path / "gate_block.yml",
        """
        version: 1
        block:
          enabled: true
          min_error_findings: 0
          critical_statuses: [GAP, INCONSISTENTE]
          critical_tokens: [show_14]
        partial:
          enabled: true
          statuses: [GAP, INCONSISTENTE]
          min_findings: 1
          banner_text: "PARCIAL"
        output:
          allow_partial_on_block: false
        """,
    )
    try:
        decision = evaluate_report_gate(dq_report, load_report_gate_policy(policy_path))
    finally:
        policy_path.unlink(missing_ok=True)

    assert decision.status == "block"
    assert decision.should_generate_output is False
    assert any("show_14_access_control" in blocker for blocker in decision.blockers)


def test_renderer_writes_manifest_with_min_fields_and_lineage(tmp_path: Path) -> None:
    """Renderer should generate report manifest with minimal schema and lineage refs."""

    dq_report = load_dq_report(_fixture_path("dq", "dq_report_gaps_inconsistencies.json"))

    template_path = tmp_path / "template_gate_lineage.docx"
    document = Document()
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "tmj_report.docx"
    result = WordReportRenderer(
        event_id=2090,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": {
                "columns": ["sessao", "presentes", "validos"],
                "rows": [{"sessao": "show_12", "presentes": 120, "validos": 140}],
            }
        },
        table_lineage_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": [
                {
                    "source_id": "SRC_ACCESS_2025",
                    "location": "page=3",
                    "evidence_text": "Tabela de controle de acesso por sessao",
                }
            ]
        },
        dq_report=dq_report,
    ).render()

    assert output_path.exists()
    assert result.manifest_path.exists()

    manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["schema_version"] == 1
    assert manifest["event_id"] == 2090
    assert manifest["template"]["path"] == str(template_path)
    assert "queries" in manifest
    assert "marts" in manifest
    assert "placeholders" in manifest
    assert "lineage_refs" in manifest
    assert manifest["dq_summary"]["error"] == 1
    assert manifest["dq_summary"]["gap_count"] == 1
    assert manifest["dq_summary"]["inconsistency_count"] == 1

    placeholders = {entry["placeholder_id"]: entry for entry in manifest["placeholders"]}
    attendance = placeholders["PUBLICO__ATTENDANCE__TABLE"]
    assert attendance["lineage_ref_ids"]
    assert attendance["replacement_count"] == 1

    lineage_by_id = {entry["lineage_ref_id"]: entry for entry in manifest["lineage_refs"]}
    first_ref = lineage_by_id[attendance["lineage_ref_ids"][0]]
    assert first_ref["source_id"] == "SRC_ACCESS_2025"
    assert first_ref["location"] == "page=3"
