"""Base DOCX renderer scaffold for TMJ report generation.

The renderer currently performs only the bootstrap behavior:
- validate template availability;
- open the template DOCX;
- optionally resolve text placeholders from mapping-driven payload;
- optionally resolve table placeholders from mapping-driven payload;
- optionally resolve figure placeholders from payload;
- append a small execution marker paragraph;
- save the generated output DOCX.

Lineage blocks are rendered for table/figure artifacts.
"""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from datetime import datetime, timezone
import logging
from pathlib import Path
from typing import Any

from docx import Document

from .placeholders_mapping import (
    DEFAULT_PLACEHOLDERS_MAPPING_PATH,
    PlaceholderRenderType,
    load_placeholders_mapping,
)
from .mart_query_runner import (
    MartQueryRunner,
    rows_to_chart_payload,
    rows_to_table_payload,
    rows_to_text_payload,
)
from .render_gaps import render_gaps_section
from .render_figure import render_figure_placeholder
from .render_lineage import extract_lineage_refs_from_rows, render_lineage_block
from .report_manifest import (
    build_report_manifest,
    default_report_manifest_path,
    write_report_manifest,
)
from .render_show_coverage import (
    SHOW_COVERAGE_PLACEHOLDER_ID,
    render_show_coverage_section,
)
from .render_table import render_table_placeholder
from .render_text import query_to_text_payload, render_text_placeholder


@dataclass(frozen=True)
class WordReportRenderResult:
    """Result contract for one scaffold render execution.

    Args:
        event_id: Event identifier used for this run.
        template_path: Source template DOCX path.
        output_path: Generated output DOCX path.
        manifest_path: Generated report manifest JSON path.
        generated_at: UTC timestamp of render completion.
    """

    event_id: int
    template_path: Path
    output_path: Path
    manifest_path: Path
    generated_at: datetime


class WordReportRenderer:
    """Render a scaffold DOCX from an existing template.

    Args:
        event_id: Event id used to identify render context.
        template_path: Path to source DOCX template.
        output_path: Path where generated DOCX will be saved.
        placeholders_mapping_path: Optional mapping YAML path.
        text_payload_by_placeholder: Optional text payload by placeholder id.
        table_payload_by_placeholder: Optional table payload by placeholder id.
        figure_payload_by_placeholder: Optional figure payload by placeholder id.
        table_lineage_by_placeholder: Optional lineage refs by table placeholder id.
        figure_lineage_by_placeholder: Optional lineage refs by figure placeholder id.
        dq_report: Optional DQ report payload to render GAP/INCONSISTENTE section.
        gaps_placeholder_id: Placeholder id used for GAP/INCONSISTENTE section.
        report_status_note: Optional status note appended in DOCX (e.g. PARCIAL).
        manifest_output_path: Optional output path for report manifest JSON.
        coverage_governance: Optional governance metadata for report manifest
            (agenda/coverage-contract versions and paths).
        mart_query_runner: Optional DB runner to auto-build payloads from mapping.
        logger: Optional logger for operational messages.

    Raises:
        ValueError: If `event_id` is not positive.
    """

    def __init__(
        self,
        *,
        event_id: int,
        template_path: Path | str,
        output_path: Path | str,
        placeholders_mapping_path: Path | str | None = None,
        text_payload_by_placeholder: Mapping[str, Any] | None = None,
        table_payload_by_placeholder: Mapping[str, Any] | None = None,
        figure_payload_by_placeholder: Mapping[str, Any] | None = None,
        table_lineage_by_placeholder: Mapping[str, Any] | None = None,
        figure_lineage_by_placeholder: Mapping[str, Any] | None = None,
        dq_report: Mapping[str, Any] | None = None,
        gaps_placeholder_id: str = "GAPS__SUMMARY__TEXT",
        report_status_note: str | None = None,
        manifest_output_path: Path | str | None = None,
        coverage_governance: Mapping[str, Any] | None = None,
        mart_query_runner: MartQueryRunner | None = None,
        logger: logging.Logger | None = None,
    ) -> None:
        if int(event_id) <= 0:
            raise ValueError("event_id deve ser inteiro positivo.")
        self.event_id = int(event_id)
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)
        if placeholders_mapping_path is None:
            self.placeholders_mapping_path = DEFAULT_PLACEHOLDERS_MAPPING_PATH
        else:
            self.placeholders_mapping_path = Path(placeholders_mapping_path)
        self.text_payload_by_placeholder = dict(text_payload_by_placeholder or {})
        self.table_payload_by_placeholder = dict(table_payload_by_placeholder or {})
        self.figure_payload_by_placeholder = dict(figure_payload_by_placeholder or {})
        self.table_lineage_by_placeholder = dict(table_lineage_by_placeholder or {})
        self.figure_lineage_by_placeholder = dict(figure_lineage_by_placeholder or {})
        self.dq_report = dict(dq_report or {}) if dq_report is not None else None
        self.gaps_placeholder_id = str(gaps_placeholder_id).strip() or "GAPS__SUMMARY__TEXT"
        status_note = str(report_status_note or "").strip()
        self.report_status_note = status_note or None
        self.manifest_output_path = (
            Path(manifest_output_path) if manifest_output_path is not None else None
        )
        self.coverage_governance = dict(coverage_governance or {})
        self.mart_query_runner = mart_query_runner
        self.logger = logger or logging.getLogger(__name__)

    def _resolve_text_payload_value(
        self,
        placeholder_id: str,
        raw_payload: Any,
    ) -> str | Sequence[str]:
        """Normalize one placeholder payload before text rendering.

        Args:
            placeholder_id: Placeholder id under processing.
            raw_payload: Raw payload from renderer input context.

        Returns:
            String or list of strings ready for `render_text_placeholder`.

        Raises:
            ValueError: If payload type is unsupported.
        """

        if isinstance(raw_payload, str):
            return raw_payload
        if isinstance(raw_payload, Sequence) and not isinstance(raw_payload, (bytes, bytearray, str)):
            if all(isinstance(item, str) for item in raw_payload):
                return list(raw_payload)
            if all(isinstance(item, Mapping) for item in raw_payload):
                return query_to_text_payload(raw_payload)
        raise ValueError(
            f"Payload invalido para placeholder {placeholder_id}: esperado str, list[str] ou list[dict]."
        )

    def _resolve_table_payload_value(
        self,
        placeholder_id: str,
        raw_payload: Any,
    ) -> tuple[list[str], list[Mapping[str, Any] | Sequence[Any]]]:
        """Normalize one table payload before table rendering.

        Args:
            placeholder_id: Placeholder id under processing.
            raw_payload: Raw payload from renderer input context.

        Returns:
            Tuple containing `(columns, rows)` ready for `render_table_placeholder`.

        Raises:
            ValueError: If payload type is unsupported or incomplete.
        """

        if not isinstance(raw_payload, Mapping):
            raise ValueError(
                f"Payload de tabela invalido para placeholder {placeholder_id}: esperado objeto."
            )

        if "columns" not in raw_payload or "rows" not in raw_payload:
            raise ValueError(
                f"Payload de tabela invalido para placeholder {placeholder_id}: "
                "campos obrigatorios columns e rows."
            )

        columns_raw = raw_payload["columns"]
        rows_raw = raw_payload["rows"]

        if not isinstance(columns_raw, Sequence) or isinstance(columns_raw, (str, bytes, bytearray)):
            raise ValueError(
                f"Payload de tabela invalido para placeholder {placeholder_id}: columns deve ser list[str]."
            )
        if not isinstance(rows_raw, Sequence) or isinstance(rows_raw, (str, bytes, bytearray)):
            raise ValueError(
                f"Payload de tabela invalido para placeholder {placeholder_id}: rows deve ser lista."
            )

        columns = list(columns_raw)
        rows = list(rows_raw)
        return columns, rows

    def _extract_query_params(self, params: Mapping[str, Any]) -> dict[str, Any]:
        """Extract SQL filter params from mapping params.

        Non-scalar params (e.g. `columns` lists for table contracts) are ignored
        for SQL filtering and kept for renderer-level adaptation only.

        Args:
            params: Placeholder mapping params.

        Returns:
            Filter params containing only scalar SQL-bindable values.
        """

        query_params: dict[str, Any] = {}
        for key, value in params.items():
            if isinstance(value, (str, int, float, bool)) or value is None:
                query_params[str(key)] = value
        return query_params

    def _resolve_figure_payload_value(
        self,
        placeholder_id: str,
        raw_payload: Any,
    ) -> tuple[Path, str, float]:
        """Normalize one figure payload before figure rendering.

        Args:
            placeholder_id: Placeholder id under processing.
            raw_payload: Raw payload from renderer input context.

        Returns:
            Tuple `(image_path, caption, width_inches)` for figure renderer.

        Raises:
            ValueError: If payload type is unsupported or incomplete.
        """

        if not isinstance(raw_payload, Mapping):
            raise ValueError(
                f"Payload de figura invalido para placeholder {placeholder_id}: esperado objeto."
            )
        image_path = raw_payload.get("image_path")
        caption = raw_payload.get("caption")
        width_inches = raw_payload.get("width_inches", 6.0)

        if not isinstance(image_path, (str, Path)) or not str(image_path).strip():
            raise ValueError(
                f"Payload de figura invalido para placeholder {placeholder_id}: image_path obrigatorio."
            )
        if not isinstance(caption, str) or not caption.strip():
            raise ValueError(
                f"Payload de figura invalido para placeholder {placeholder_id}: caption obrigatorio."
            )

        try:
            width_value = float(width_inches)
        except (TypeError, ValueError) as exc:
            raise ValueError(
                f"Payload de figura invalido para placeholder {placeholder_id}: width_inches invalido."
            ) from exc

        return Path(image_path), caption.strip(), width_value

    def render(self) -> WordReportRenderResult:
        """Generate the scaffold output DOCX from template.

        Returns:
            `WordReportRenderResult` with output path and timestamp.

        Raises:
            FileNotFoundError: If template path does not exist.
            ValueError: If template path is not `.docx`.
            ValueError: If payload placeholder is unknown or invalid.
            OSError: If output file cannot be written.
        """

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template DOCX nao encontrado: {self.template_path}")
        if self.template_path.suffix.lower() != ".docx":
            raise ValueError("template_path deve apontar para arquivo .docx valido.")

        self.logger.info(
            "Gerando scaffold de relatorio Word: event_id=%s template=%s output=%s",
            self.event_id,
            self.template_path,
            self.output_path,
        )

        document = Document(self.template_path)
        mapping = load_placeholders_mapping(self.placeholders_mapping_path)
        replacement_count_by_placeholder: dict[str, int] = {}
        mart_result_by_placeholder: dict[str, Any] = {}
        if self.dq_report is not None and self.gaps_placeholder_id not in self.text_payload_by_placeholder:
            self.text_payload_by_placeholder[self.gaps_placeholder_id] = render_gaps_section(
                self.dq_report
            )

        if self.mart_query_runner is not None:
            chart_payload_count = 0
            for item in mapping.placeholders:
                if item.render_type == PlaceholderRenderType.TEXT:
                    if item.placeholder_id in self.text_payload_by_placeholder:
                        continue
                    result = self.mart_query_runner.execute_view(
                        item.mart_name,
                        params=self._extract_query_params(item.params),
                        section_cache_key=item.placeholder_id,
                    )
                    mart_result_by_placeholder[item.placeholder_id] = result
                    self.text_payload_by_placeholder[item.placeholder_id] = rows_to_text_payload(
                        result.rows
                    )
                elif item.render_type == PlaceholderRenderType.TABLE:
                    if item.placeholder_id in self.table_payload_by_placeholder:
                        continue
                    result = self.mart_query_runner.execute_view(
                        item.mart_name,
                        params=self._extract_query_params(item.params),
                        section_cache_key=item.placeholder_id,
                    )
                    mart_result_by_placeholder[item.placeholder_id] = result
                    columns = item.params.get("columns")
                    columns_list = columns if isinstance(columns, list) else None
                    self.table_payload_by_placeholder[item.placeholder_id] = rows_to_table_payload(
                        result.rows,
                        columns=columns_list,
                    )
                    if item.placeholder_id not in self.table_lineage_by_placeholder:
                        self.table_lineage_by_placeholder[item.placeholder_id] = (
                            extract_lineage_refs_from_rows(result.rows)
                        )
                elif item.render_type == PlaceholderRenderType.FIGURE:
                    # Figure rendering is implemented in a dedicated issue. We still
                    # query and shape payload here to keep DB contract validation early.
                    result = self.mart_query_runner.execute_view(
                        item.mart_name,
                        params=self._extract_query_params(item.params),
                        section_cache_key=item.placeholder_id,
                    )
                    mart_result_by_placeholder[item.placeholder_id] = result
                    rows_to_chart_payload(result.rows)
                    if item.placeholder_id not in self.figure_lineage_by_placeholder:
                        self.figure_lineage_by_placeholder[item.placeholder_id] = (
                            extract_lineage_refs_from_rows(result.rows)
                        )
                    chart_payload_count += 1
            if chart_payload_count:
                self.logger.info(
                    "Payload(s) de figura preparados via mart_query_runner: %s",
                    chart_payload_count,
                )

        text_placeholder_ids = {
            item.placeholder_id
            for item in mapping.placeholders
            if item.render_type == PlaceholderRenderType.TEXT
        }
        table_placeholder_ids = {
            item.placeholder_id
            for item in mapping.placeholders
            if item.render_type == PlaceholderRenderType.TABLE
        }
        figure_placeholder_ids = {
            item.placeholder_id
            for item in mapping.placeholders
            if item.render_type == PlaceholderRenderType.FIGURE
        }

        unknown_payload_ids = sorted(
            placeholder_id
            for placeholder_id in self.text_payload_by_placeholder
            if placeholder_id not in text_placeholder_ids
        )
        if unknown_payload_ids:
            unknown_str = ", ".join(unknown_payload_ids)
            raise ValueError(
                "Payload de texto contem placeholder(s) nao mapeado(s) como text: "
                f"{unknown_str}."
            )
        unknown_table_payload_ids = sorted(
            placeholder_id
            for placeholder_id in self.table_payload_by_placeholder
            if placeholder_id not in table_placeholder_ids
        )
        if unknown_table_payload_ids:
            unknown_str = ", ".join(unknown_table_payload_ids)
            raise ValueError(
                "Payload de tabela contem placeholder(s) nao mapeado(s) como table: "
                f"{unknown_str}."
            )
        unknown_figure_payload_ids = sorted(
            placeholder_id
            for placeholder_id in self.figure_payload_by_placeholder
            if placeholder_id not in figure_placeholder_ids
        )
        if unknown_figure_payload_ids:
            unknown_str = ", ".join(unknown_figure_payload_ids)
            raise ValueError(
                "Payload de figura contem placeholder(s) nao mapeado(s) como figure: "
                f"{unknown_str}."
            )

        for placeholder_id, raw_payload in self.text_payload_by_placeholder.items():
            value = self._resolve_text_payload_value(placeholder_id, raw_payload)
            replacements = render_text_placeholder(document, placeholder_id, value)
            replacement_count_by_placeholder[placeholder_id] = replacements
            self.logger.info(
                "Placeholder de texto renderizado: %s | ocorrencias=%s",
                placeholder_id,
                replacements,
            )
        for placeholder_id, raw_payload in self.table_payload_by_placeholder.items():
            columns, rows = self._resolve_table_payload_value(placeholder_id, raw_payload)
            if placeholder_id == SHOW_COVERAGE_PLACEHOLDER_ID:
                replacements = render_show_coverage_section(
                    document,
                    placeholder_id=placeholder_id,
                    columns=columns,
                    rows=rows,
                )
            else:
                replacements = render_table_placeholder(
                    document,
                    placeholder_id,
                    columns=columns,
                    rows=rows,
                )
            replacement_count_by_placeholder[placeholder_id] = replacements
            self.logger.info(
                "Placeholder de tabela renderizado: %s | ocorrencias=%s",
                placeholder_id,
                replacements,
            )
            lineage_lines = render_lineage_block(
                self.table_lineage_by_placeholder.get(placeholder_id),
                placeholder_id=placeholder_id,
                artifact_type="table",
            )
            for line in lineage_lines:
                document.add_paragraph(line)
        for placeholder_id, raw_payload in self.figure_payload_by_placeholder.items():
            image_path, caption, width_inches = self._resolve_figure_payload_value(
                placeholder_id,
                raw_payload,
            )
            replacements = render_figure_placeholder(
                document,
                placeholder_id,
                image_path,
                caption,
                width_inches=width_inches,
            )
            replacement_count_by_placeholder[placeholder_id] = replacements
            self.logger.info(
                "Placeholder de figura renderizado: %s | ocorrencias=%s",
                placeholder_id,
                replacements,
            )
            lineage_lines = render_lineage_block(
                self.figure_lineage_by_placeholder.get(placeholder_id),
                placeholder_id=placeholder_id,
                artifact_type="figure",
            )
            for line in lineage_lines:
                document.add_paragraph(line)

        generated_at = datetime.now(timezone.utc)
        if self.report_status_note:
            document.add_paragraph(f"[STATUS] {self.report_status_note}")
        document.add_paragraph(
            (
                "[NPBB scaffold] Report generated automatically. "
                f"event_id={self.event_id}; generated_at={generated_at.isoformat()}"
            )
        )

        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(self.output_path)
        self.logger.info("DOCX scaffold gerado com sucesso em %s", self.output_path)

        lineage_by_placeholder: dict[str, Any] = {}
        lineage_by_placeholder.update(self.table_lineage_by_placeholder)
        lineage_by_placeholder.update(self.figure_lineage_by_placeholder)
        for placeholder_id, mart_result in mart_result_by_placeholder.items():
            if placeholder_id in lineage_by_placeholder:
                continue
            rows = getattr(mart_result, "rows", None)
            if not isinstance(rows, Sequence):
                continue
            extracted_refs = extract_lineage_refs_from_rows(rows)
            if extracted_refs:
                lineage_by_placeholder[placeholder_id] = extracted_refs

        manifest = build_report_manifest(
            event_id=self.event_id,
            template_path=self.template_path,
            output_docx_path=self.output_path,
            generated_at=generated_at,
            mapping=mapping,
            placeholder_replacements=replacement_count_by_placeholder,
            placeholder_lineage=lineage_by_placeholder,
            mart_results=mart_result_by_placeholder,
            dq_report=self.dq_report,
            coverage_governance=self.coverage_governance,
        )
        manifest_path = (
            self.manifest_output_path
            if self.manifest_output_path is not None
            else default_report_manifest_path(self.output_path)
        )
        final_manifest_path = write_report_manifest(manifest, manifest_path)
        self.logger.info("Manifest do relatorio gerado em %s", final_manifest_path)

        return WordReportRenderResult(
            event_id=self.event_id,
            template_path=self.template_path,
            output_path=self.output_path,
            manifest_path=final_manifest_path,
            generated_at=generated_at,
        )
