"""Unit tests for show-coverage Word section renderer."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from reports.word import (
    SHOW_COVERAGE_PLACEHOLDER_ID,
    ShowCoverageRenderError,
    render_show_coverage_section,
)
from reports.word.render_text import placeholder_token


def _build_template(path: Path, placeholder_id: str) -> Path:
    """Create a DOCX template with one placeholder paragraph."""

    document = Document()
    document.add_paragraph(placeholder_token(placeholder_id))
    document.save(path)
    return path


def test_render_show_coverage_section_renders_table_and_legend(tmp_path: Path) -> None:
    """Show-coverage renderer should fill table and append status legend lines."""

    template_path = _build_template(tmp_path / "show_coverage_template.docx", SHOW_COVERAGE_PLACEHOLDER_ID)
    doc = Document(template_path)

    replaced = render_show_coverage_section(
        doc,
        columns=[
            "dia",
            "sessao",
            "status_access_control",
            "status_optin",
            "status_ticket_sales",
            "observacoes",
        ],
        rows=[
            {
                "dia": "2025-12-12",
                "sessao": "Show 12/12",
                "status_access_control": "ok",
                "status_optin": "partial",
                "status_ticket_sales": "gap",
                "observacoes": "falta_vendidos_total",
            }
        ],
    )
    out_path = tmp_path / "show_coverage_out.docx"
    doc.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    assert len(rendered.tables) == 1
    header = [cell.text for cell in rendered.tables[0].rows[0].cells]
    assert header == [
        "dia",
        "sessao",
        "status_access_control",
        "status_optin",
        "status_ticket_sales",
        "observacoes",
    ]
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert placeholder_token(SHOW_COVERAGE_PLACEHOLDER_ID) not in all_text
    assert "Legenda de status da cobertura de shows por dia:" in all_text
    assert "OK: dataset presente e carregado para a sessao." in all_text
    assert "GAP: dataset ausente para a sessao esperada." in all_text
    assert "INCONSISTENTE: dataset presente com conflito entre fontes ou regras." in all_text


def test_render_show_coverage_section_requires_required_columns(tmp_path: Path) -> None:
    """Show-coverage renderer should fail when required columns are missing."""

    template_path = _build_template(tmp_path / "show_coverage_invalid.docx", SHOW_COVERAGE_PLACEHOLDER_ID)
    doc = Document(template_path)

    with pytest.raises(ShowCoverageRenderError, match="faltam colunas obrigatorias"):
        render_show_coverage_section(
            doc,
            columns=[
                "dia",
                "sessao",
                "status_access_control",
                "status_ticket_sales",
                "observacoes",
            ],
            rows=[],
        )
