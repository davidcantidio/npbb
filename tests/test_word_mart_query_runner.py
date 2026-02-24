"""Unit tests for Word mart query runner and payload adapters."""

from __future__ import annotations

from pathlib import Path

import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from reports.word.mart_query_runner import (
    MartQueryRunner,
    MartQueryRunnerError,
    rows_to_chart_payload,
    rows_to_table_payload,
    rows_to_text_payload,
)
from reports.word.renderer import WordReportRenderer


def _create_sqlite_session(tmp_path: Path) -> Session:
    """Create SQLite session with deterministic mart views for tests."""

    db_path = tmp_path / "marts_runner.db"
    engine = create_engine(f"sqlite:///{db_path}")
    with engine.begin() as connection:
        connection.exec_driver_sql("DROP VIEW IF EXISTS mart_report_sources")
        connection.exec_driver_sql("DROP VIEW IF EXISTS mart_report_attendance_by_session")
        connection.exec_driver_sql(
            """
            CREATE VIEW mart_report_sources AS
            SELECT
              'Fonte A com cobertura de acesso' AS text,
              'ok' AS status
            """
        )
        connection.exec_driver_sql(
            """
            CREATE VIEW mart_report_attendance_by_session AS
            SELECT
              'show_12' AS sessao,
              120 AS presentes,
              150 AS validos
            UNION ALL
            SELECT
              'show_13' AS sessao,
              180 AS presentes,
              200 AS validos
            """
        )
    return Session(engine)


def _write_mapping_yaml(path: Path) -> Path:
    """Write temporary mapping with text and table placeholders."""

    path.write_text(
        "\n".join(
            [
                "version: 1",
                "placeholders:",
                "  - placeholder_id: FONTES__SUMMARY__TEXT",
                "    mart_name: mart_report_sources",
                "    params: {}",
                "    render_type: text",
                "  - placeholder_id: PUBLICO__ATTENDANCE__TABLE",
                "    mart_name: mart_report_attendance_by_session",
                "    params:",
                "      columns:",
                "        - sessao",
                "        - presentes",
                "        - validos",
                "    render_type: table",
            ]
        ),
        encoding="utf-8",
    )
    return path


def test_mart_query_runner_executes_view_and_uses_cache(tmp_path: Path) -> None:
    """Runner should execute mart and reuse cached result by section key."""

    with _create_sqlite_session(tmp_path) as session:
        runner = MartQueryRunner(session, enable_cache=True)

        first = runner.execute_view(
            "mart_report_sources",
            section_cache_key="FONTES__SUMMARY__TEXT",
        )
        second = runner.execute_view(
            "mart_report_sources",
            section_cache_key="FONTES__SUMMARY__TEXT",
        )

        assert first.cached is False
        assert second.cached is True
        assert first.rows == second.rows
        assert first.rows[0]["text"] == "Fonte A com cobertura de acesso"


def test_mart_query_runner_reports_missing_view_with_actionable_error(tmp_path: Path) -> None:
    """Runner should fail with mart name when view does not exist."""

    with _create_sqlite_session(tmp_path) as session:
        runner = MartQueryRunner(session)
        with pytest.raises(MartQueryRunnerError, match="mart_report_inexistente"):
            runner.execute_view("mart_report_inexistente")


def test_rows_adapters_return_typed_payloads() -> None:
    """Adapters should shape rows into text/table/chart payload contracts."""

    rows = [
        {"dia": "2025-12-12", "presentes": 120, "text": "Cobertura inicial"},
        {"dia": "2025-12-13", "presentes": 180, "text": "Cobertura expandida"},
    ]

    text_payload = rows_to_text_payload(rows)
    table_payload = rows_to_table_payload(rows, columns=["dia", "presentes"])
    chart_payload = rows_to_chart_payload(rows, x_field="dia", y_field="presentes")

    assert text_payload == ["Cobertura inicial", "Cobertura expandida"]
    assert table_payload["columns"] == ["dia", "presentes"]
    assert table_payload["rows"][0]["dia"] == "2025-12-12"
    assert chart_payload["x_field"] == "dia"
    assert chart_payload["y_field"] == "presentes"
    assert len(chart_payload["data"]) == 2


def test_renderer_integrates_mart_query_runner_with_mapping(tmp_path: Path) -> None:
    """Renderer should auto-build text/table payloads from mapping via mart runner."""

    template_path = tmp_path / "template_runner.docx"
    from docx import Document  # local import to keep test dependency scoped

    document = Document()
    document.add_paragraph("{{FONTES__SUMMARY__TEXT}}")
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.save(template_path)

    mapping_path = _write_mapping_yaml(tmp_path / "mapping.yml")
    output_path = tmp_path / "out" / "report_runner.docx"

    with _create_sqlite_session(tmp_path) as session:
        runner = MartQueryRunner(session, enable_cache=True)
        renderer = WordReportRenderer(
            event_id=2040,
            template_path=template_path,
            output_path=output_path,
            placeholders_mapping_path=mapping_path,
            mart_query_runner=runner,
        )
        renderer.render()

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "{{FONTES__SUMMARY__TEXT}}" not in text
    assert "{{PUBLICO__ATTENDANCE__TABLE}}" not in text
    assert "Fonte A com cobertura de acesso" in text
    assert len(rendered.tables) == 1
    assert [cell.text for cell in rendered.tables[0].rows[0].cells] == [
        "sessao",
        "presentes",
        "validos",
    ]


def test_renderer_with_runner_fails_for_missing_mart_view(tmp_path: Path) -> None:
    """Renderer should fail early when mapping references non-existent mart view."""

    template_path = tmp_path / "template_missing_view.docx"
    from docx import Document

    document = Document()
    document.add_paragraph("{{FONTES__SUMMARY__TEXT}}")
    document.save(template_path)

    mapping_path = tmp_path / "mapping_missing.yml"
    mapping_path.write_text(
        "\n".join(
            [
                "version: 1",
                "placeholders:",
                "  - placeholder_id: FONTES__SUMMARY__TEXT",
                "    mart_name: mart_report_nao_existe",
                "    params: {}",
                "    render_type: text",
            ]
        ),
        encoding="utf-8",
    )

    with _create_sqlite_session(tmp_path) as session:
        runner = MartQueryRunner(session)
        renderer = WordReportRenderer(
            event_id=2041,
            template_path=template_path,
            output_path=tmp_path / "out" / "report_missing_view.docx",
            placeholders_mapping_path=mapping_path,
            mart_query_runner=runner,
        )
        with pytest.raises(ValueError, match="mart_report_nao_existe"):
            renderer.render()
