"""Unit tests for Word table placeholder rendering helpers."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from reports.word import (
    PlaceholderTableRenderError,
    placeholder_token,
    render_table_placeholder,
)


def _build_doc_with_paragraph(path: Path, text: str) -> Document:
    """Create DOCX document with a single paragraph text and save it."""

    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return document


def test_render_table_placeholder_renders_headers_and_rows(tmp_path: Path) -> None:
    """Table renderer should replace token and create table with expected order."""

    placeholder_id = "PUBLICO__ATTENDANCE__TABLE"
    template_path = tmp_path / "table_template.docx"
    _build_doc_with_paragraph(template_path, placeholder_token(placeholder_id))
    document = Document(template_path)

    replaced = render_table_placeholder(
        document,
        placeholder_id,
        columns=["sessao", "presentes", "validos"],
        rows=[
            {"sessao": "show_12", "presentes": 100, "validos": 120},
            {"sessao": "show_13", "presentes": 150, "validos": 200},
        ],
    )
    out_path = tmp_path / "table_out.docx"
    document.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert placeholder_token(placeholder_id) not in text
    assert len(rendered.tables) == 1

    table = rendered.tables[0]
    header = [cell.text for cell in table.rows[0].cells]
    assert header == ["sessao", "presentes", "validos"]
    assert [cell.text for cell in table.rows[1].cells] == ["show_12", "100", "120"]
    assert [cell.text for cell in table.rows[2].cells] == ["show_13", "150", "200"]


def test_render_table_placeholder_raises_for_missing_required_column(tmp_path: Path) -> None:
    """Renderer should fail when one mart row does not satisfy expected columns."""

    placeholder_id = "PUBLICO__ATTENDANCE__TABLE"
    template_path = tmp_path / "table_missing_column.docx"
    _build_doc_with_paragraph(template_path, placeholder_token(placeholder_id))
    document = Document(template_path)

    with pytest.raises(PlaceholderTableRenderError, match="Contrato de colunas violado"):
        render_table_placeholder(
            document,
            placeholder_id,
            columns=["sessao", "presentes", "validos"],
            rows=[{"sessao": "show_12", "presentes": 100}],
        )


def test_render_table_placeholder_raises_for_missing_token(tmp_path: Path) -> None:
    """Renderer should fail with actionable error when placeholder token is absent."""

    template_path = tmp_path / "table_missing_placeholder.docx"
    _build_doc_with_paragraph(template_path, "Sem placeholder aqui.")
    document = Document(template_path)

    with pytest.raises(PlaceholderTableRenderError, match="nao encontrado"):
        render_table_placeholder(
            document,
            "PUBLICO__ATTENDANCE__TABLE",
            columns=["sessao"],
            rows=[],
        )
