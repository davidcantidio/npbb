"""Template placeholder validation for Word report generation.

This module inspects DOCX templates and validates placeholder coverage
against the configured mapping contract.
"""

from __future__ import annotations

from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from .placeholders_mapping import WordPlaceholdersMapping

_PLACEHOLDER_TOKEN_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


@dataclass(frozen=True)
class TemplatePlaceholderFinding:
    """One placeholder validation finding.

    Args:
        code: Finding code identifier.
        placeholder_id: Placeholder identifier detected in template.
        message: Human-readable and actionable message.
    """

    code: str
    placeholder_id: str
    message: str


def _iter_document_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    """Iterate all paragraphs in body, tables, headers and footers.

    Args:
        doc: Open DOCX document.

    Yields:
        Paragraph nodes from all supported document containers.
    """

    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def find_placeholders(docx_path: Path | str) -> set[str]:
    """Extract placeholder identifiers from one DOCX file.

    Supported token format:
    - `{{PLACEHOLDER_ID}}`
    - `{{ PLACEHOLDER_ID }}` (spaces are normalized)

    Args:
        docx_path: Path to source DOCX file.

    Returns:
        Set of unique placeholder identifiers found in the document.

    Raises:
        FileNotFoundError: If DOCX path does not exist.
        ValueError: If file extension is not `.docx`.
    """

    template_path = Path(docx_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template DOCX nao encontrado: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError("docx_path deve apontar para arquivo .docx valido.")

    document = Document(template_path)
    placeholders: set[str] = set()
    for paragraph in _iter_document_paragraphs(document):
        for match in _PLACEHOLDER_TOKEN_RE.findall(paragraph.text):
            placeholder = str(match).strip()
            if placeholder:
                placeholders.add(placeholder)
    return placeholders


def validate_template_placeholders(
    placeholders: set[str],
    mapping: WordPlaceholdersMapping,
) -> list[TemplatePlaceholderFinding]:
    """Validate template placeholders against mapping contract.

    Args:
        placeholders: Placeholders detected in template.
        mapping: Loaded placeholder mapping contract.

    Returns:
        Findings list. Empty list means validation succeeded.
    """

    mapped_ids = {item.placeholder_id for item in mapping.placeholders}
    unknown_ids = sorted(placeholders - mapped_ids)

    findings: list[TemplatePlaceholderFinding] = []
    for placeholder_id in unknown_ids:
        findings.append(
            TemplatePlaceholderFinding(
                code="UNKNOWN_PLACEHOLDER",
                placeholder_id=placeholder_id,
                message=(
                    "Placeholder presente no template sem mapping configurado. "
                    "Como corrigir: adicionar item em word_placeholders.yml."
                ),
            )
        )
    return findings
