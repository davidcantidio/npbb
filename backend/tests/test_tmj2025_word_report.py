from __future__ import annotations

from datetime import date, datetime, timezone
from pathlib import Path

from docx import Document
from sqlalchemy.pool import StaticPool
from sqlmodel import SQLModel, Session, create_engine

from app.models.models import (
    AttendanceAccessControl,
    BbRelationshipSegment,
    EventSessionType,
    FestivalLead,
    OptinTransaction,
    Source,
    SourceKind,
    TicketCategorySegmentMap,
)
from app.reports.tmj2025_word import RenderOptions, generate_tmj2025_closing_report
from app.services.tmj_segments import normalize_ticket_category
from app.services.tmj_sessions import get_or_create_tmj2025_session


def make_engine():
    return create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )


def _make_min_template(path: Path) -> None:
    doc = Document()

    # Title-like placeholder (not used by renderer).
    doc.add_heading("Fechamento TMJ 2025 (Template v2)", level=0)

    def h1(text: str) -> None:
        doc.add_heading(text, level=1)
        doc.add_paragraph("PLACEHOLDER")

    def h2(text: str) -> None:
        doc.add_heading(text, level=2)
        doc.add_paragraph("PLACEHOLDER")

    h1("1. Contexto do evento")
    h1("2. Objetivo do relatorio")
    h1("3. Fontes de dados e limitacoes")

    h1("4. Big numbers (recorte analisado)")
    h2("4.1 Publico do evento (controle de acesso - entradas validadas)")
    h2("4.2 Dinamica de vendas (pre-venda) - shows (Opt-in aceitos)")
    h2("4.3 Quem sao clientes do Banco (proxy via categoria de ingresso - Opt-in)")

    h1("5. Perfil do publico (DIMAC - 12 a 14/12)")
    h2("5.1 Satisfacao e percepcao (DIMAC)")

    h1("6. Pre-venda (leitura tecnica e aprendizados)")
    h1("7. Performance nas redes (Instagram e social listening)")
    h1("8. Midia e imprensa (MTC)")
    h1("9. Leads e ativacoes (Festival de Esportes | 12-14/12)")
    h1("10. Recomendacoes (2026) - acoes tecnicas e de produto")
    h1("11. Apendice - definicoes rapidas")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def test_generate_tmj2025_closing_report_inserts_content_and_sources(tmp_path: Path) -> None:
    engine = make_engine()
    SQLModel.metadata.create_all(engine)

    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "out.docx"
    _make_min_template(template_path)

    with Session(engine) as session:
        # Register sources referenced by facts.
        session.add(
            Source(
                source_id="SRC_XLSX_OPTIN_ACEITOS_DOZE",
                kind=SourceKind.XLSX,
                uri=str(tmp_path / "optin.xlsx"),
                display_name="optin.xlsx",
            )
        )
        session.add(
            Source(
                source_id="SRC_PDF_ACESSO_NOTURNO_TREZE",
                kind=SourceKind.PDF,
                uri=str(tmp_path / "acesso.pdf"),
                display_name="acesso.pdf",
            )
        )
        session.add(
            Source(
                source_id="SRC_XLSX_LEADS_FESTIVAL_ESPORTES",
                kind=SourceKind.XLSX,
                uri=str(tmp_path / "leads.xlsx"),
                display_name="leads.xlsx",
            )
        )
        session.commit()

        # Create expected show sessions for 12/13/14 with sparse coverage.
        s12 = get_or_create_tmj2025_session(
            session,
            event_id=None,
            session_date=date(2025, 12, 12),
            session_type=EventSessionType.NOTURNO_SHOW,
            source_of_truth_source_id="SRC_XLSX_OPTIN_ACEITOS_DOZE",
        )
        s13 = get_or_create_tmj2025_session(
            session,
            event_id=None,
            session_date=date(2025, 12, 13),
            session_type=EventSessionType.NOTURNO_SHOW,
            source_of_truth_source_id="SRC_PDF_ACESSO_NOTURNO_TREZE",
        )
        s14 = get_or_create_tmj2025_session(
            session,
            event_id=None,
            session_date=date(2025, 12, 14),
            session_type=EventSessionType.NOTURNO_SHOW,
            source_of_truth_source_id=None,
        )

        # Opt-in facts for 12 and 13 (none for 14).
        cat_raw = "FUNCIONARIOS BB"
        cat_norm = normalize_ticket_category(cat_raw)
        session.add(
            TicketCategorySegmentMap(
                ticket_category_raw=cat_raw,
                ticket_category_norm=cat_norm,
                segment=BbRelationshipSegment.FUNCIONARIO_BB,
                inferred=True,
                inference_rule="test",
            )
        )
        session.add(
            OptinTransaction(
                session_id=int(s12.id),
                source_id="SRC_XLSX_OPTIN_ACEITOS_DOZE",
                sheet_name="01.1 - Opt-In Aceitos",
                row_number=1,
                purchase_at=datetime(2025, 10, 14, 9, 0, 0, tzinfo=timezone.utc),
                purchase_date=date(2025, 10, 14),
                ticket_category_raw=cat_raw,
                ticket_category_norm=cat_norm,
                ticket_qty=2,
                person_key_hash="a" * 64,
            )
        )
        session.add(
            OptinTransaction(
                session_id=int(s13.id),
                source_id="SRC_XLSX_OPTIN_ACEITOS_DOZE",
                sheet_name="01.1 - Opt-In Aceitos",
                row_number=2,
                purchase_at=datetime(2025, 10, 15, 9, 0, 0, tzinfo=timezone.utc),
                purchase_date=date(2025, 10, 15),
                ticket_category_raw=cat_raw,
                ticket_category_norm=cat_norm,
                ticket_qty=1,
                person_key_hash="b" * 64,
            )
        )

        # Access control for show 13 only.
        session.add(
            AttendanceAccessControl(
                session_id=int(s13.id),
                source_id="SRC_PDF_ACESSO_NOTURNO_TREZE",
                ingressos_validos=100,
                presentes=90,
                ausentes=10,
                pdf_page=1,
                evidence="Tabela resumo",
            )
        )

        # Leads on 14/12.
        session.add(
            FestivalLead(
                event_id=None,
                session_id=None,
                source_id="SRC_XLSX_LEADS_FESTIVAL_ESPORTES",
                sheet_name="Entidades",
                row_number=1,
                lead_created_at=datetime(2025, 12, 14, 14, 0, 0, tzinfo=timezone.utc),
                lead_created_date=date(2025, 12, 14),
                person_key_hash="c" * 64,
                acoes="BB :: CARAMELO",
            )
        )
        session.commit()

        generate_tmj2025_closing_report(
            session,
            template_path=template_path,
            output_path=output_path,
            options=RenderOptions(session_key_prefix="TMJ2025_"),
        )

    out_doc = Document(str(output_path))
    texts = [p.text for p in out_doc.paragraphs if p.text]

    assert "PLACEHOLDER" not in " ".join(texts)
    assert any("Resumo agregado do recorte carregado no banco" in t for t in texts)
    assert any(t.startswith("Fonte:") for t in texts)
    assert any("Shows por dia (auditoria de cobertura minima)" in t for t in texts)
    assert out_doc.tables, "expected at least one table in output docx"

