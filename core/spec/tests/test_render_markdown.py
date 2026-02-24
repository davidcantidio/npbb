"""Unit tests for DOCX-as-spec markdown rendering."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from core.spec.render_markdown import build_initial_checklist_rows, render_docx_as_spec_md
from core.spec.spec_models import (
    CHECKLIST_COLUMNS,
    MANDATORY_SHOWS_BY_DAY_ITEM,
    DocxSpec,
    DocxSection,
)
from etl.cli_spec import main


def _build_spec(*, sections: list[DocxSection]) -> DocxSpec:
    """Create a minimal DocxSpec fixture used by render tests."""
    return DocxSpec(
        source_path="template.docx",
        extracted_at=datetime.now(timezone.utc),
        sections=sections,
        figures=[],
        tables=[],
    )


def test_render_markdown_includes_fixed_checklist_columns_and_default_status() -> None:
    """Checklist markdown uses fixed columns and starts rows as GAP."""
    spec = _build_spec(
        sections=[
            DocxSection(level=1, title="1. Resumo 2025", body_lines=[], start_paragraph_index=0),
            DocxSection(level=2, title="Indicadores de Publico", body_lines=[], start_paragraph_index=3),
        ]
    )

    rendered = render_docx_as_spec_md(spec)

    assert "| " + " | ".join(CHECKLIST_COLUMNS) + " |" in rendered
    assert "| Resumo 2025 | TODO | TODO | TODO | Secao \"Resumo 2025\" (template DOCX) | GAP |" in rendered
    assert MANDATORY_SHOWS_BY_DAY_ITEM in rendered


def test_build_initial_checklist_rows_avoids_duplicate_mandatory_show_item() -> None:
    """Mandatory show coverage row is added only once."""
    spec = _build_spec(
        sections=[
            DocxSection(level=1, title=MANDATORY_SHOWS_BY_DAY_ITEM, body_lines=[], start_paragraph_index=0),
        ]
    )

    rows = build_initial_checklist_rows(spec)
    matching = [row for row in rows if row.section == MANDATORY_SHOWS_BY_DAY_ITEM]

    assert len(matching) == 1


def test_cli_spec_extract_writes_markdown_with_mandatory_show_item(tmp_path: Path) -> None:
    """`spec:extract` emits markdown checklist including required show-coverage row."""
    docx_path = tmp_path / "template.docx"
    output_md = tmp_path / "00_docx_as_spec.md"

    doc = Document()
    doc.add_heading("Resumo Executivo", level=1)
    doc.add_paragraph("Texto base.")
    doc.save(docx_path)

    rc = main(
        [
            "spec:extract",
            "--docx",
            str(docx_path),
            "--out",
            str(output_md),
            "--format",
            "md",
        ]
    )

    assert rc == 0
    assert output_md.exists()
    content = output_md.read_text(encoding="utf-8")
    assert MANDATORY_SHOWS_BY_DAY_ITEM in content
