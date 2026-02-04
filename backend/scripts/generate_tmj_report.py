"""Gera relatorio TMJ 2025 em DOCX (agregado, sem PII)."""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from sqlmodel import Session

from app.db.database import engine
from app.db.metadata import SQLModel
from app.schemas.dashboard_leads_report import DashboardLeadsReportQuery
from app.services.dashboard_leads_report import build_dashboard_leads_report


def _parse_date(value: str | None) -> date | None:
    if not value:
        return None
    return date.fromisoformat(value)


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Metrica"
    hdr[1].text = "Valor"
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value


def generate_docx(output_path: Path, report) -> None:
    doc = Document()

    doc.add_heading("Relatorio Tecnico - Festival TMJ 2025", level=0)
    doc.add_paragraph(f"Gerado em: {report.generated_at.isoformat()}")

    doc.add_heading("1. Contexto e objetivo", level=1)
    doc.add_paragraph(
        "Este relatorio consolida dados agregados do Festival TMJ 2025 para apoiar decisoes "
        "de marketing, ativacao e melhoria de captação."
    )

    doc.add_heading("2. Big Numbers", level=1)
    _add_kv_table(
        doc,
        [
            ("Total de leads", str(report.big_numbers.total_leads)),
            ("Total de compras", str(report.big_numbers.total_compras)),
            ("Publico do evento", str(report.big_numbers.total_publico)),
            ("Taxa de conversao (%)", f"{report.big_numbers.taxa_conversao:.2f}"),
            ("Criterio publico", report.big_numbers.criterio_publico),
            ("Criterio compras", report.big_numbers.criterio_compras),
        ],
    )

    doc.add_heading("3. Publico do evento", level=1)
    doc.add_paragraph(
        "A definicao de publico segue o criterio informado no bloco de big numbers. "
        "Se o evento nao possui publico_realizado/projetado, usamos leads como proxy."
    )

    doc.add_heading("4. Perfil do publico", level=1)
    doc.add_paragraph("Distribuicao de idade (percentual sobre registros com idade).")
    idade_rows = [
        (item.faixa, f"{item.total} ({item.percentual:.2f}%)")
        for item in report.perfil_publico.distribuicao_idade
    ]
    _add_kv_table(doc, idade_rows or [("Sem dados", "0")])
    doc.add_paragraph(f"% sem idade: {report.perfil_publico.percent_sem_idade:.2f}%")

    doc.add_paragraph("Distribuicao de genero (percentual sobre total).")
    genero_rows = [
        (item.faixa, f"{item.total} ({item.percentual:.2f}%)")
        for item in report.perfil_publico.distribuicao_genero
    ]
    _add_kv_table(doc, genero_rows or [("Sem dados", "0")])
    doc.add_paragraph(f"% sem genero: {report.perfil_publico.percent_sem_genero:.2f}%")

    doc.add_heading("5. Clientes BB", level=1)
    doc.add_paragraph(
        f"Total clientes BB: {report.clientes_bb.total_clientes_bb} "
        f"({report.clientes_bb.percentual_clientes_bb:.2f}%). "
        f"Criterio: {report.clientes_bb.criterio_usado}."
    )

    doc.add_heading("6. Pre-venda", level=1)
    doc.add_paragraph(
        f"Janela pre-venda: {report.pre_venda.janela_pre_venda or 'nao definida'}"
    )
    doc.add_paragraph(
        f"Volume pre-venda: {report.pre_venda.volume_pre_venda} | "
        f"Volume geral: {report.pre_venda.volume_venda_geral}"
    )
    if report.pre_venda.observacao:
        doc.add_paragraph(report.pre_venda.observacao)

    doc.add_heading("7. Performance em redes", level=1)
    doc.add_paragraph(report.redes.observacao or "Sem dados de redes.")

    doc.add_heading("8. Limitacoes e dados faltantes", level=1)
    if report.dados_faltantes:
        for item in report.dados_faltantes:
            doc.add_paragraph(f"- {item}")
    else:
        doc.add_paragraph("Sem lacunas identificadas.")

    doc.add_heading("9. Recomendacoes tecnicas (top 5)", level=1)
    recommendations = [
        "Adicionar indicador de cliente BB no cadastro/import de leads.",
        "Integrar dados de redes sociais ao data lake do evento.",
        "Garantir evento_id nos leads de ticketing para filtros consistentes.",
        "Padronizar coleta de genero/idade para reduzir percentual sem dado.",
        "Automatizar checagens de qualidade do import (colunas obrigatorias).",
    ]
    for rec in recommendations:
        doc.add_paragraph(f"- {rec}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Gerar relatorio TMJ 2025 em DOCX.")
    parser.add_argument("--evento-id", type=int, default=None, help="ID do evento (preferencial).")
    parser.add_argument(
        "--evento-nome",
        type=str,
        default="Festival TMJ 2025",
        help="Nome do evento (fallback).",
    )
    parser.add_argument("--data-inicio", type=str, default=None, help="Data inicio (YYYY-MM-DD).")
    parser.add_argument("--data-fim", type=str, default=None, help="Data fim (YYYY-MM-DD).")
    parser.add_argument(
        "--output",
        type=str,
        default=str(Path(__file__).resolve().parents[2] / "reports" / "Festival_TMJ_2025_relatorio.docx"),
        help="Caminho do arquivo DOCX.",
    )
    args = parser.parse_args()

    params = DashboardLeadsReportQuery(
        data_inicio=_parse_date(args.data_inicio),
        data_fim=_parse_date(args.data_fim),
        evento_id=args.evento_id,
        evento_nome=args.evento_nome,
    )

    if engine.url.drivername.startswith("sqlite"):
        SQLModel.metadata.create_all(engine)

    with Session(engine) as session:
        report = build_dashboard_leads_report(session, params)

    output_path = Path(args.output)
    generate_docx(output_path, report)
    print(f"Relatorio gerado em: {output_path}")


if __name__ == "__main__":
    main()
