"""Unit tests for Word text placeholder rendering helpers."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from reports.word import (
    PlaceholderTextRenderError,
    placeholder_token,
    query_to_text_payload,
    render_text_placeholder,
)


def _build_doc_with_paragraph(path: Path, text: str) -> Document:
    """Create DOCX document with a single paragraph text and save it."""

    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return document


def test_render_text_placeholder_replaces_scalar_text(tmp_path: Path) -> None:
    """Scalar payload should replace placeholder token in paragraph text."""

    placeholder_id = "FONTES__SUMMARY__TEXT"
    template_path = tmp_path / "scalar.docx"
    _build_doc_with_paragraph(template_path, f"Resumo: {placeholder_token(placeholder_id)}")
    document = Document(template_path)

    replaced = render_text_placeholder(document, placeholder_id, "Fonte A e Fonte B com limitacoes.")
    out_path = tmp_path / "scalar_out.docx"
    document.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert placeholder_token(placeholder_id) not in all_text
    assert "Resumo: Fonte A e Fonte B com limitacoes." in all_text


def test_render_text_placeholder_renders_bullets_for_list_payload(tmp_path: Path) -> None:
    """List payload should render bullet-like paragraphs when token is standalone."""

    placeholder_id = "FONTES__SUMMARY__TEXT"
    template_path = tmp_path / "bullets.docx"
    _build_doc_with_paragraph(template_path, placeholder_token(placeholder_id))
    document = Document(template_path)

    replaced = render_text_placeholder(
        document,
        placeholder_id,
        ["Fonte A com status ok", "Fonte B com GAP de vendas"],
    )
    out_path = tmp_path / "bullets_out.docx"
    document.save(out_path)

    assert replaced == 1
    rendered = Document(out_path)
    paragraphs = [paragraph.text for paragraph in rendered.paragraphs if paragraph.text.strip()]
    assert "Fonte A com status ok" in paragraphs
    assert "Fonte B com GAP de vendas" in paragraphs


def test_render_text_placeholder_raises_for_missing_token(tmp_path: Path) -> None:
    """Renderer should fail with actionable error when placeholder token is absent."""

    template_path = tmp_path / "missing.docx"
    _build_doc_with_paragraph(template_path, "Sem placeholder aqui.")
    document = Document(template_path)

    with pytest.raises(PlaceholderTextRenderError, match="nao encontrado"):
        render_text_placeholder(document, "FONTES__SUMMARY__TEXT", "Texto")


def test_query_to_text_payload_adapts_rows_to_short_strings() -> None:
    """Adapter should convert mart rows to concise layout-ready strings."""

    lines = query_to_text_payload(
        [
            {"text": "Topline de alcance em redes"},
            {"label": "fontes_com_gap", "value": 2},
            {"metric": "share_bb", "value": 0.42, "session": "show_12"},
        ]
    )
    assert lines[0] == "Topline de alcance em redes"
    assert "fontes_com_gap: 2" in lines
    assert any("metric=share_bb" in line for line in lines)
