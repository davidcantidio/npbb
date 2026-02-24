"""Unit tests for DOCX figure/table inventory extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from core.spec.docx_figures_tables import (
    extract_figures,
    extract_figures_from_docx,
    extract_tables,
    extract_tables_from_docx,
)
from core.spec.docx_sections import DocxSectionExtractor, find_nearest_section_title


def _build_sample_docx(tmp_path: Path) -> Path:
    """Create a small DOCX fixture with headings, figure captions and one table."""
    path = tmp_path / "sample_spec.docx"
    doc = Document()
    doc.add_heading("Resumo Executivo", level=1)
    doc.add_paragraph("Texto introdutorio.")
    doc.add_paragraph("Figura 1 - Entradas validadas por dia")
    doc.add_heading("Shows", level=1)
    doc.add_paragraph("Figura 2 - Publico por dia de show")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Metrica"
    table.rows[0].cells[1].text = "12/12"
    table.rows[0].cells[2].text = "13/12"
    table.rows[1].cells[0].text = "Entradas validadas"
    table.rows[1].cells[1].text = "1000"
    table.rows[1].cells[2].text = "1200"
    doc.save(path)
    return path


def test_extract_figures_associates_to_nearest_section(tmp_path: Path) -> None:
    """Figure inventory keeps caption text, position, and nearest section."""
    docx_path = _build_sample_docx(tmp_path)
    sections = DocxSectionExtractor().extract_sections(docx_path)
    doc = Document(docx_path)

    figures = extract_figures(doc.paragraphs, sections)

    assert len(figures) == 2
    assert figures[0].caption == "Figura 1 - Entradas validadas por dia"
    assert figures[0].paragraph_index == 2
    assert figures[0].section_title == "Resumo Executivo"
    assert figures[1].caption == "Figura 2 - Publico por dia de show"
    assert figures[1].paragraph_index == 4
    assert figures[1].section_title == "Shows"
    assert find_nearest_section_title(4, sections) == "Shows"


def test_extract_tables_captures_header_shape_and_section(tmp_path: Path) -> None:
    """Table inventory captures headers, dimensions, and closest heading context."""
    docx_path = _build_sample_docx(tmp_path)
    doc = Document(docx_path)

    tables = extract_tables(doc)

    assert len(tables) == 1
    assert tables[0].table_index == 1
    assert tables[0].n_rows == 2
    assert tables[0].n_cols == 3
    assert tables[0].header_cells == ["Metrica", "12/12", "13/12"]
    assert tables[0].section_title == "Shows"


def test_docx_path_wrappers_delegate_to_core_extractors(tmp_path: Path) -> None:
    """Path-based wrappers return the same inventories as in-memory extractors."""
    docx_path = _build_sample_docx(tmp_path)
    sections = DocxSectionExtractor().extract_sections(docx_path)
    doc = Document(docx_path)

    expected_figures = extract_figures(doc.paragraphs, sections)
    expected_tables = extract_tables(doc)

    assert extract_figures_from_docx(docx_path, sections) == expected_figures
    assert extract_tables_from_docx(docx_path) == expected_tables
