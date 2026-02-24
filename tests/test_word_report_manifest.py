"""Unit tests for report manifest generation in Word report renderer."""

from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace
from datetime import datetime, timezone

from docx import Document

from reports.word import (
    PlaceholderRenderType,
    WordPlaceholderBinding,
    WordPlaceholdersMapping,
    WordReportRenderer,
    build_report_manifest,
)


def test_renderer_generates_manifest_with_lineage_and_dq_summary(tmp_path: Path) -> None:
    """Renderer should persist report_manifest.json with required audit fields."""

    template_path = tmp_path / "template_manifest.docx"
    document = Document()
    document.add_paragraph("{{PUBLICO__ATTENDANCE__TABLE}}")
    document.add_paragraph("{{GAPS__SUMMARY__TEXT}}")
    document.save(template_path)

    output_path = tmp_path / "out" / "report_manifest.docx"
    dq_report = {
        "summary": {"error": 1, "warning": 2, "info": 0},
        "sections": {
            "gaps_by_dataset": {
                "show_14_access_control": [
                    {"message": "Sem fonte oficial para show de 14/12"}
                ]
            },
            "inconsistencies": [],
        },
    }
    renderer = WordReportRenderer(
        event_id=2040,
        template_path=template_path,
        output_path=output_path,
        table_payload_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": {
                "columns": ["sessao", "presentes", "validos"],
                "rows": [{"sessao": "show_12", "presentes": 100, "validos": 120}],
            }
        },
        table_lineage_by_placeholder={
            "PUBLICO__ATTENDANCE__TABLE": [
                {
                    "source_id": "SRC_ACCESS_12",
                    "location": "page=3",
                    "evidence_text": "Tabela de controle de acesso",
                }
            ]
        },
        dq_report=dq_report,
    )
    result = renderer.render()

    expected_manifest_path = output_path.parent / "report_manifest.json"
    assert result.manifest_path == expected_manifest_path
    assert expected_manifest_path.exists()

    manifest = json.loads(expected_manifest_path.read_text(encoding="utf-8"))
    assert manifest["event_id"] == 2040
    assert manifest["template"]["path"] == str(template_path)
    assert manifest["dq_summary"]["error"] == 1
    assert manifest["dq_summary"]["gap_count"] == 1

    placeholders = {item["placeholder_id"]: item for item in manifest["placeholders"]}
    attendance = placeholders["PUBLICO__ATTENDANCE__TABLE"]
    assert attendance["mart_name"] == "mart_report_attendance_by_session"
    assert attendance["lineage_ref_ids"]

    lineage_by_id = {item["lineage_ref_id"]: item for item in manifest["lineage_refs"]}
    lineage_ref = lineage_by_id[attendance["lineage_ref_ids"][0]]
    assert lineage_ref["source_id"] == "SRC_ACCESS_12"
    assert lineage_ref["location"] == "page=3"
    assert manifest["sources"] == ["SRC_ACCESS_12"]


def test_build_report_manifest_includes_query_metadata(tmp_path: Path) -> None:
    """Manifest builder should include query execution metadata by placeholder."""

    template_path = tmp_path / "template_version_anchor.docx"
    template_path.write_bytes(b"placeholder")
    output_docx_path = tmp_path / "out" / "report.docx"

    mapping = WordPlaceholdersMapping(
        version=3,
        placeholders=(
            WordPlaceholderBinding(
                placeholder_id="FONTES__SUMMARY__TEXT",
                mart_name="mart_report_sources",
                params={"include_notes": True},
                render_type=PlaceholderRenderType.TEXT,
                spec_item="Fontes e limitacoes",
                description=None,
            ),
        ),
    )
    manifest = build_report_manifest(
        event_id=2041,
        template_path=template_path,
        output_docx_path=output_docx_path,
        generated_at=result_time(),
        mapping=mapping,
        placeholder_replacements={"FONTES__SUMMARY__TEXT": 1},
        placeholder_lineage={
            "FONTES__SUMMARY__TEXT": [
                {
                    "source_id": "SRC_SUMMARY",
                    "location": "sheet=fontes|range=A1:B2",
                    "evidence_text": "Tabela de fontes",
                }
            ]
        },
        mart_results={
            "FONTES__SUMMARY__TEXT": SimpleNamespace(
                rows=(
                    {
                        "source_id": "SRC_SUMMARY",
                        "location": "sheet=fontes|range=A1:B2",
                        "evidence_text": "Tabela de fontes",
                    },
                ),
                cached=True,
            )
        },
        dq_report=None,
    )

    assert manifest["mapping_version"] == 3
    assert manifest["queries"][0]["executed"] is True
    assert manifest["queries"][0]["row_count"] == 1
    assert manifest["queries"][0]["cached"] is True
    assert manifest["marts"][0]["mart_name"] == "mart_report_sources"
    assert manifest["marts"][0]["query_count"] == 1


def result_time() -> datetime:
    """Return current UTC timestamp string for manifest builder fixture."""

    return datetime.now(timezone.utc)
