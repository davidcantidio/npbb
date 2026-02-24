"""Tests for Word text rendering and template placeholder validation."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

from reports.word import WordReportRenderer
from reports.word.placeholders_mapping import load_placeholders_mapping
from reports.word.template_validate import (
    find_placeholders,
    validate_template_placeholders,
)


_TOKEN_RE = re.compile(r"\{\{\s*[^{}]+?\s*\}\}")


def _fixture_template_path() -> Path:
    """Return path to the minimal Word report template fixture."""

    return Path(__file__).resolve().parent / "fixtures" / "docx" / "report_template_min.docx"


def test_word_report_fixture_has_known_text_placeholder() -> None:
    """Fixture must expose mapped text placeholder for renderer coverage."""

    placeholders = find_placeholders(_fixture_template_path())
    assert placeholders == {"FONTES__SUMMARY__TEXT"}


def test_validate_template_placeholders_returns_unknown_when_unmapped(tmp_path: Path) -> None:
    """Validator should flag placeholder ids that do not exist in mapping."""

    template_path = tmp_path / "template_unknown.docx"
    document = Document(_fixture_template_path())
    document.add_paragraph("{{INEXISTENTE__BLOCO__TEXT}}")
    document.save(template_path)

    placeholders = find_placeholders(template_path)
    mapping = load_placeholders_mapping()
    findings = validate_template_placeholders(placeholders, mapping)

    assert len(findings) == 1
    assert findings[0].code == "UNKNOWN_PLACEHOLDER"
    assert findings[0].placeholder_id == "INEXISTENTE__BLOCO__TEXT"


def test_word_report_renderer_replaces_text_placeholder_and_removes_tokens(tmp_path: Path) -> None:
    """Renderer should replace mapped text placeholder and remove token from output."""

    output_path = tmp_path / "report_rendered.docx"
    renderer = WordReportRenderer(
        event_id=2025,
        template_path=_fixture_template_path(),
        output_path=output_path,
        text_payload_by_placeholder={
            "FONTES__SUMMARY__TEXT": [
                "Controle de acesso com evidencias por pagina.",
                "Opt-in de pre-venda auditado por aba e range.",
            ]
        },
    )

    renderer.render()

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Controle de acesso com evidencias por pagina." in text
    assert "Opt-in de pre-venda auditado por aba e range." in text
    assert _TOKEN_RE.search(text) is None
